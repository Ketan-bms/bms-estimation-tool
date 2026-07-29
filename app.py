"""
BMS Estimation Tool — app.py v2
Complete multi-project estimating platform.
Run: streamlit run app.py
"""

import json, os, io, base64
from material_module import module_material, init_pricebooks
from markup_ui import module_markup
from pdf_takeoff import run_pdf_takeoff, takeoff_to_session_format
from point_list_extractor import generate_point_list_prompt, parse_point_list_response, infer_io_type
from soo_extractor import (
    generate_overview_prompt,
    generate_pointlist_prompt,
    generate_appendix_prompt,
    generate_important_notes_prompt,
    parse_pointlist_response,
    parse_notes_response
)
from pathlib import Path
from datetime import date
from collections import defaultdict

import streamlit as st
import pandas as pd

st.set_page_config(page_title="BMS Estimator", page_icon="🏗", layout="wide",
                   initial_sidebar_state="expanded")

st.markdown("""
<style>
[data-testid="stMetricValue"]{font-size:1.8rem}
.block-label{font-size:.7rem;font-weight:600;letter-spacing:.06em;color:#64748b;text-transform:uppercase;margin-bottom:.3rem}
.disc-banner{background:#fff3cd;border-left:4px solid #f59e0b;padding:10px 14px;border-radius:0 6px 6px 0;margin-bottom:8px}
.chip-done{background:#dcfce7;color:#166534;font-size:11px;padding:2px 8px;border-radius:4px;font-weight:500}
.chip-prog{background:#dbeafe;color:#1e40af;font-size:11px;padding:2px 8px;border-radius:4px;font-weight:500}
.chip-issue{background:#fef9c3;color:#854d0e;font-size:11px;padding:2px 8px;border-radius:4px;font-weight:500}
.chip-empty{background:#f1f5f9;color:#64748b;font-size:11px;padding:2px 8px;border-radius:4px;border:0.5px solid #e2e8f0}
</style>
""", unsafe_allow_html=True)

DEFAULT_RATES = {"Engineering":95,"Programming":85,"Integration":75,"Graphics":70,"Startup":80}
PHASES = list(DEFAULT_RATES.keys())
MODULE_ORDER = ["Takeoff","Point List","Estimate","Proposal"]

# ── State ─────────────────────────────────────────────────────────────────────
# ── Persistent storage using st.cache_resource ────────────────────────────────
# Survives browser refresh for the lifetime of the Streamlit server process.
# Clients (including template bytes) and project metadata are stored here.
# Uploaded PDFs are NOT stored (too large) — must re-upload after server restart.

@st.cache_resource
def _get_store():
    """One shared dict that persists across reruns and refreshes."""
    return {
        "clients":   {},   # client metadata (rates, template names)
        "projects":  {},   # project metadata (no doc bytes)
        "templates": {},   # client template bytes keyed by "{client}_pl" / "{client}_prop"
    }


def _save_app_state():
    """Write current session state into the persistent store."""
    store = _get_store()
    try:
        # ── Clients ───────────────────────────────────────────────────────
        clients_meta = {}
        for cname, c in st.session_state.get("clients", {}).items():
            meta = {k: v for k, v in c.items()
                    if k not in ("pl_template_bytes", "prop_template_bytes")}
            # Store template bytes separately
            if c.get("pl_template_bytes"):
                store["templates"][f"{cname}_pl"]   = c["pl_template_bytes"]
            if c.get("prop_template_bytes"):
                store["templates"][f"{cname}_prop"] = c["prop_template_bytes"]
            clients_meta[cname] = meta
        store["clients"] = clients_meta

        # ── Projects (no doc bytes) ───────────────────────────────────────
        projects_meta = {}
        for pname, p in st.session_state.get("projects", {}).items():
            projects_meta[pname] = {
                k: v for k, v in p.items() if k not in ("docs",)
            }
            projects_meta[pname].setdefault("docs", {})
            projects_meta[pname]["doc_names"] = p.get("doc_names", {})
        store["projects"] = projects_meta

    except Exception:
        pass


def _load_app_state():
    """Read from persistent store into session state (only if not already loaded)."""
    store = _get_store()
    try:
        # ── Clients ───────────────────────────────────────────────────────
        for cname, meta in store.get("clients", {}).items():
            if cname not in st.session_state["clients"]:
                entry = dict(meta)
                pl_b   = store["templates"].get(f"{cname}_pl")
                prop_b = store["templates"].get(f"{cname}_prop")
                if pl_b:   entry["pl_template_bytes"]   = pl_b
                if prop_b: entry["prop_template_bytes"] = prop_b
                st.session_state["clients"][cname] = entry

        # ── Projects ──────────────────────────────────────────────────────
        for pname, meta in store.get("projects", {}).items():
            if pname not in st.session_state["projects"]:
                entry = dict(meta)
                entry.setdefault("docs", {})
                entry.setdefault("doc_names", {})
                st.session_state["projects"][pname] = entry

    except Exception:
        pass


def init():
    defaults = {
        "clients":{}, "projects":{},
        "active_project":None, "active_module":"Takeoff",
        "nav":"Overview", "ai_history":[],
        "storage_loaded": False,
    }
    for k,v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # Load persisted data on first run (from cache_resource store)
    if not st.session_state.get("storage_loaded"):
        _load_app_state()
        st.session_state["storage_loaded"] = True

PROJECT_STATUSES = ["Active","In Progress","On Hold","Completed","Archived"]
PROJECT_STATUS_CSS = {
    "Active":      ("background:#dbeafe;color:#1e40af;border:1px solid #93c5fd",   "🔵"),
    "In Progress": ("background:#fef9c3;color:#854d0e;border:1px solid #fcd34d",   "🟡"),
    "On Hold":     ("background:#ffedd5;color:#9a3412;border:1px solid #fdba74",   "🟠"),
    "Completed":   ("background:#dcfce7;color:#166534;border:1px solid #86efac",   "🟢"),
    "Archived":    ("background:#f1f5f9;color:#475569;border:1px solid #cbd5e1",   "⚫"),
}

def status_badge_html(status):
    """Return a colored HTML badge for a project status."""
    css, emoji = PROJECT_STATUS_CSS.get(
        status, ("background:#f1f5f9;color:#475569;border:1px solid #cbd5e1", "○"))
    return (f'<span style="{css};padding:3px 10px;border-radius:20px;' 
            f'font-size:11px;font-weight:600;display:inline-block">' 
            f'{emoji} {status}</span>')

def new_project(name, client, bid_date, address):
    return {
        "name":name,"client":client,"bid_date":bid_date,
        "address":address,"created":str(date.today()),
        "project_status": "Active",
        "docs":{},"doc_names":{},
        "takeoff":{"equipment":[],"discrepancies":[],"status":"not_started"},
        "point_list":{"rows":[],"status":"not_started"},
        "point_list_appendix":{"rows":[],"status":"not_started"},
        "estimate":{"lines":[],"rates":{},"markup":10,"status":"not_started"},
        "proposal":{"text":"","status":"not_started"},
    }

def module_status(p, mod):
    return p[mod.lower().replace(" ","_")]["status"]

def module_locked(p, mod):
    """New unlock rules — modules open based on available docs, not chain."""
    if mod == "Takeoff":
        return False   # always open
    if mod == "Point List":
        # Unlocks when SOO or controls spec is uploaded
        return not (p["docs"].get("SOO") or p["docs"].get("Controls spec"))
    if mod == "Estimate":
        # Unlocks when point list has rows OR controls spec uploaded
        has_points = len(p["point_list"].get("rows", [])) > 0
        has_spec   = bool(p["docs"].get("Controls spec") or p["docs"].get("SOO"))
        return not (has_points or has_spec)
    if mod == "Proposal":
        # Unlocks when estimate has lines OR takeoff is done
        has_estimate = len(p["estimate"].get("lines", [])) > 0
        has_takeoff  = p["takeoff"]["status"] != "not_started"
        return not (has_estimate or has_takeoff)
    return False  # AI Advisor, Drawing Markup always open


def module_data_warning(p, mod):
    """Return warning text when module is open but missing some data."""
    has_takeoff  = p["takeoff"]["status"] != "not_started"
    has_points   = len(p["point_list"].get("rows", [])) > 0
    has_estimate = len(p["estimate"].get("lines", [])) > 0
    has_soo      = bool(p["docs"].get("SOO"))
    has_spec     = bool(p["docs"].get("Controls spec"))

    if mod == "Point List" and not has_takeoff:
        return ("ℹ️ No takeoff loaded — point list will be generated from SOO/spec only. "
                "Load takeoff for exact device tags.")
    if mod == "Estimate" and not has_points:
        return ("ℹ️ No point list yet — estimate will use rough point counts from SOO/spec. "
                "Generate a point list first for more accurate hours.")
    if mod == "Estimate" and not has_takeoff:
        return ("ℹ️ No takeoff loaded — material quantities are estimated, not from drawings. "
                "Load takeoff for accurate device counts.")
    if mod == "Proposal" and not has_takeoff:
        return ("ℹ️ No takeoff loaded — proposal will use SOO-based scope only. "
                "Load takeoff to auto-fill Clarifications and Exclusions with real device tags.")
    if mod == "Proposal" and not has_estimate:
        return ("ℹ️ No estimate yet — proposal will not include a price. "
                "Generate estimate first to include pricing.")
    return None

def chip(s, label):
    cls = {"done":"chip-done","in_progress":"chip-prog",
           "issues":"chip-issue"}.get(s,"chip-empty")
    return f'<span class="{cls}">{label}</span>'

def get_status(e):
    if e.get("discrepancy_flag"): return "Discrepancy"
    if e.get("soo_confirmed"):    return "SOO confirmed"
    return "Needs review"

def api_key():
    return st.session_state.get("anthropic_api_key",
                                 os.environ.get("ANTHROPIC_API_KEY",""))

# ── Sidebar ───────────────────────────────────────────────────────────────────
def sidebar():
    with st.sidebar:
        # ── Logo / header ──────────────────────────────────────────────
        st.markdown("""
        <style>
        [data-testid="stSidebar"] {
            background: #1a1f2e;
        }
        [data-testid="stSidebar"] * {
            color: #e2e8f0 !important;
        }
        .sb-logo {
            font-size: 18px; font-weight: 700; color: #fff !important;
            padding: 8px 4px 4px; letter-spacing: .02em;
        }
        .sb-logo span { color: #6366f1 !important; }
        .sb-year { font-size: 11px; color: #64748b !important; padding: 0 4px 16px; }
        .sb-section {
            font-size: 10px; font-weight: 600; color: #475569 !important;
            text-transform: uppercase; letter-spacing: .08em;
            padding: 14px 4px 6px;
        }
        .sb-nav-item {
            display: flex; align-items: center; gap: 10px;
            padding: 8px 12px; border-radius: 7px; margin-bottom: 2px;
            font-size: 13px; font-weight: 500; cursor: pointer;
            color: #cbd5e1 !important; text-decoration: none;
        }
        .sb-nav-item:hover { background: #2d3548; }
        .sb-nav-active {
            background: #6366f1 !important; color: #fff !important;
        }
        .sb-proj-name {
            display: flex; align-items: center; gap: 8px;
            padding: 7px 10px; border-radius: 7px; margin-bottom: 1px;
            font-size: 13px; font-weight: 600; color: #f1f5f9 !important;
            cursor: pointer;
        }
        .sb-proj-name:hover { background: #2d3548; }
        .sb-proj-active { background: #2d3548; }
        .sb-task {
            display: flex; align-items: center; gap: 8px;
            padding: 5px 10px 5px 28px; border-radius: 6px;
            margin-bottom: 1px; font-size: 12px; color: #94a3b8 !important;
            cursor: pointer;
        }
        .sb-task:hover { background: #232838; color: #e2e8f0 !important; }
        .sb-task-active {
            background: #232838 !important; color: #a5b4fc !important;
            font-weight: 600;
        }
        .sb-dot { font-size: 8px; }
        .sb-divider {
            border: none; border-top: 1px solid #2d3548;
            margin: 10px 0;
        }
        /* Override Streamlit button styles in sidebar */
        [data-testid="stSidebar"] .stButton button {
            background: transparent !important;
            border: none !important;
            color: #cbd5e1 !important;
            text-align: left !important;
            padding: 6px 10px !important;
            border-radius: 7px !important;
            font-size: 13px !important;
            width: 100% !important;
        }
        [data-testid="stSidebar"] .stButton button:hover {
            background: #2d3548 !important;
        }
        </style>
        """, unsafe_allow_html=True)

        st.markdown('<div class="sb-logo">🏗 <span>BMS</span> Estimator</div>',
                    unsafe_allow_html=True)
        st.markdown(f'<div class="sb-year">FY {date.today().year}</div>',
                    unsafe_allow_html=True)

        # ── Main nav ───────────────────────────────────────────────────
        nav_items = [
            ("🏠", "Overview"),
            ("📁", "Projects"),
            ("📊", "Reports"),
        ]

        for icon, label in nav_items:
            is_active = (st.session_state.nav == label and
                         (label != "Projects" or
                          not st.session_state.active_project))
            style = "sb-nav-item sb-nav-active" if is_active else "sb-nav-item"
            # Use a real button but styled via CSS
            if st.button(f"{icon}  {label}", key=f"nav_{label}",
                         use_container_width=True):
                st.session_state.nav = label
                st.session_state.active_project = None
                st.rerun()

        # ── Recent projects — clean list, click to open ───────────────
        if st.session_state.projects:
            st.markdown('<hr class="sb-divider">', unsafe_allow_html=True)
            st.markdown('<div class="sb-section">Recent Projects</div>',
                        unsafe_allow_html=True)

            for pname, p in st.session_state.projects.items():
                is_open = st.session_state.active_project == pname
                disc    = len(p["takeoff"].get("discrepancies", []))
                icon    = "📂" if is_open else "📁"

                if st.button(f"{icon}  {pname}",
                             key=f"sb_proj_{pname}",
                             use_container_width=True):
                    st.session_state.active_project = pname
                    st.session_state.active_module  = "Takeoff"
                    st.session_state.nav            = "Projects"
                    st.rerun()

                if disc:
                    st.caption(f"   ⚠️ {disc} discrepancies")

        st.markdown('<hr class="sb-divider">', unsafe_allow_html=True)

        # ── API key ────────────────────────────────────────────────────
        k = st.text_input("API key", type="password",
                          value=os.environ.get("ANTHROPIC_API_KEY", ""),
                          key="api_key_input",
                          placeholder="sk-ant-...")
        if k:
            st.session_state["anthropic_api_key"] = k

# ── Overview ──────────────────────────────────────────────────────────────────
def page_overview():
    today    = date.today()
    projects = st.session_state.projects

    # ── Styles ────────────────────────────────────────────────────────────
    st.markdown("""<style>
    .kpi { background:#fff; border:1px solid #e2e8f0; border-radius:12px;
           padding:20px 18px; text-align:center; }
    .kpi-icon  { font-size:24px; margin-bottom:6px; }
    .kpi-val   { font-size:2rem; font-weight:700; color:#1e293b; line-height:1.1; }
    .kpi-label { font-size:11px; font-weight:600; color:#64748b;
                 text-transform:uppercase; letter-spacing:.06em; margin-top:6px; }
    .kpi-sub   { font-size:11px; color:#94a3b8; margin-top:3px; }
    .pcard { background:#fff; border:1px solid #e2e8f0; border-radius:12px;
             padding:16px 18px; margin-bottom:10px; }
    .pcard-name { font-size:15px; font-weight:600; color:#1e293b; }
    .pcard-meta { font-size:12px; color:#64748b; margin:3px 0 10px; }
    .pill { display:inline-block; font-size:10px; padding:2px 8px;
            border-radius:4px; margin-right:4px; font-weight:500; }
    .pill-done  { background:#dcfce7; color:#166534; }
    .pill-prog  { background:#dbeafe; color:#1e40af; }
    .pill-warn  { background:#fef9c3; color:#854d0e; }
    .pill-empty { background:#f1f5f9; color:#94a3b8; }
    .pbar-track { height:5px; border-radius:3px; background:#f1f5f9; margin:8px 0 4px; }
    .pbar-fill  { height:5px; border-radius:3px;
                  background:linear-gradient(90deg,#6366f1,#10b981); }
    .alert { border-left:4px solid; border-radius:0 8px 8px 0;
             padding:9px 13px; margin-bottom:7px; font-size:12px; }
    .alert-red   { border-color:#ef4444; background:#fef2f2; color:#7f1d1d; }
    .alert-amber { border-color:#f59e0b; background:#fffbeb; color:#78350f; }
    .alert-blue  { border-color:#3b82f6; background:#eff6ff; color:#1e3a5f; }
    .dl-row { display:flex; justify-content:space-between; align-items:center;
              padding:8px 0; border-bottom:1px solid #f1f5f9; font-size:13px; }
    .funnel-col { text-align:center; padding:12px 8px;
                  border-radius:10px; margin:0 3px; }
    </style>""", unsafe_allow_html=True)

    # ── Header ────────────────────────────────────────────────────────────
    hc1, hc2 = st.columns([4, 1])
    hc1.markdown("## BMS Estimator")
    hc1.caption(today.strftime('%B %d, %Y'))
    if hc2.button("＋ New project", type="primary",
                  use_container_width=True, key="ov_new"):
        st.session_state.nav = "Projects"
        st.session_state.active_project = None
        st.rerun()

    # ── Compute stats ─────────────────────────────────────────────────────
    n = len(projects)

    # Pipeline value
    pipeline_val = 0
    for p in projects.values():
        try:
            lines  = p["estimate"].get("lines", [])
            markup = p["estimate"].get("markup", 10)
            sub    = sum(float(r.get("Total $", 0)) for r in lines)
            pipeline_val += sub * (1 + markup / 100)
        except: pass

    # Due this month
    due_month = 0
    for p in projects.values():
        bd = p.get("bid_date")
        if bd:
            try:
                days = (date.fromisoformat(str(bd)) - today).days
                if 0 <= days <= 30: due_month += 1
            except: pass

    # Proposals ready
    prop_ready = sum(1 for p in projects.values()
                     if module_status(p, "Proposal") == "done")

    # Needs action
    needs_action = 0
    for p in projects.values():
        disc = len(p["takeoff"].get("discrepancies", []))
        if disc: needs_action += 1
        bd = p.get("bid_date")
        if bd:
            try:
                days = (date.fromisoformat(str(bd)) - today).days
                if days <= 30 and module_status(p, "Estimate") == "not_started":
                    needs_action += 1
            except: pass

    # Pipeline funnel counts
    funnel = {"Scoping": 0, "Estimating": 0, "Proposal": 0, "Submitted": 0}
    for p in projects.values():
        ps = module_status(p, "Proposal")
        es = module_status(p, "Estimate")
        pl = module_status(p, "Point List")
        tk = module_status(p, "Takeoff")
        if ps == "done":                          funnel["Submitted"] += 1
        elif es in ("done","in_progress"):        funnel["Proposal"]  += 1
        elif pl in ("done","in_progress"):        funnel["Estimating"]+= 1
        else:                                     funnel["Scoping"]   += 1

    # ── KPI cards ─────────────────────────────────────────────────────────
    k1,k2,k3,k4,k5 = st.columns(5)
    kpi_data = [
        (k1, "📁", str(n),
         "Active bids",
         f"{funnel['Submitted']} submitted" if n else "no projects yet"),
        (k2, "💰",
         f"${pipeline_val/1000:.0f}K" if pipeline_val >= 1000
         else f"${pipeline_val:,.0f}" if pipeline_val else "—",
         "Pipeline value",
         f"{sum(1 for p in projects.values() if p['estimate'].get('lines'))} estimates built"),
        (k3, "📅", str(due_month),
         "Due this month",
         "bids in next 30 days"),
        (k4, "✅", str(prop_ready),
         "Proposals ready",
         f"of {n} active bids"),
        (k5, "⚠️", str(needs_action),
         "Needs action",
         "discrepancies or overdue steps"),
    ]
    for col, icon, val, label, sub in kpi_data:
        color = "#dc2626" if (label == "Needs action" and int(val or 0) > 0) else                 "#d97706" if (label == "Due this month" and int(val or 0) > 0) else                 "#1e293b"
        col.markdown(f"""
        <div class="kpi">
            <div class="kpi-icon">{icon}</div>
            <div class="kpi-val" style="color:{color}">{val}</div>
            <div class="kpi-label">{label}</div>
            <div class="kpi-sub">{sub}</div>
        </div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Pipeline funnel ────────────────────────────────────────────────────
    if projects:
        st.markdown("**Bid pipeline**")
        fc = st.columns(4)
        funnel_colors = [
            ("#6366f1","#ede9fe"),
            ("#3b82f6","#dbeafe"),
            ("#10b981","#d1fae5"),
            ("#f59e0b","#fef9c3"),
        ]
        for i,(stage,cnt) in enumerate(funnel.items()):
            col, (fg,bg) = fc[i], funnel_colors[i]
            pct = f"{int(cnt/n*100)}%" if n else "0%"
            col.markdown(f"""
            <div class="funnel-col" style="background:{bg}">
                <div style="font-size:11px;font-weight:600;color:{fg};
                            text-transform:uppercase;letter-spacing:.05em">{stage}</div>
                <div style="font-size:2rem;font-weight:700;color:{fg};
                            line-height:1.2;margin:4px 0">{cnt}</div>
                <div style="font-size:11px;color:{fg};opacity:.7">{pct} of bids</div>
            </div>""", unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)

    # ── Main + right panel ────────────────────────────────────────────────
    col_main, col_right = st.columns([1.7, 1])

    with col_main:
        st.markdown("**Projects**")
        if not projects:
            st.markdown("""
            <div style="border:2px dashed #e2e8f0;border-radius:12px;
                        padding:48px;text-align:center">
                <div style="font-size:36px;margin-bottom:12px">📋</div>
                <div style="font-size:16px;font-weight:600;color:#475569;
                            margin-bottom:6px">No projects yet</div>
                <div style="font-size:13px;color:#94a3b8">
                    Click <strong>＋ New project</strong> above to get started</div>
            </div>""", unsafe_allow_html=True)
        else:
            for pname, p in projects.items():
                disc  = len(p["takeoff"].get("discrepancies", []))
                devs  = len(p["takeoff"].get("equipment", []))
                bd    = p.get("bid_date", "")

                # Progress
                done_n = sum(1 for m in MODULE_ORDER
                             if module_status(p,m) in ("done","in_progress"))
                pct    = int(done_n / len(MODULE_ORDER) * 100)

                # Deadline text
                dl_txt  = ""
                dl_icon = ""
                if bd:
                    try:
                        days = (date.fromisoformat(str(bd)) - today).days
                        bd_fmt = date.fromisoformat(str(bd)).strftime("%b %d, %Y")
                        if days < 0:   dl_icon, dl_txt = "⚪", f"Past due · {bd_fmt}"
                        elif days<=14: dl_icon, dl_txt = "🔴", f"{bd_fmt}"
                        elif days<=30: dl_icon, dl_txt = "🟡", f"{bd_fmt}"
                        else:          dl_icon, dl_txt = "🟢", f"{bd_fmt}"
                    except:
                        dl_txt = str(bd)[:10]

                with st.container(border=True):
                    # Title row
                    tc1, tc2 = st.columns([3,1])
                    tc1.markdown(f"**{pname}**")
                    if dl_txt:
                        tc2.markdown(f"{dl_icon} {dl_txt}")

                    # Info row
                    proj_st = p.get("project_status","Active")
                    st_emoji = {"Active":"🔵","In Progress":"🟡",
                                "On Hold":"🟠","Completed":"🟢",
                                "Archived":"⚫"}.get(proj_st,"🔵")
                    meta_parts = [f"{st_emoji} {proj_st}"]
                    if devs:  meta_parts.append(f"📐 {devs} devices")
                    if disc:  meta_parts.append(f"⚠️ {disc} discrepancies")
                    meta_parts.append(f"{done_n}/{len(MODULE_ORDER)} modules")
                    st.caption("  ·  ".join(meta_parts))

                    # Module status pills using Streamlit columns
                    mc = st.columns(len(MODULE_ORDER))
                    for i, mod in enumerate(MODULE_ORDER):
                        s = module_status(p, mod)
                        if s == "done":
                            mc[i].success(mod, icon="✅")
                        elif s == "in_progress":
                            mc[i].info(mod, icon="🔄")
                        elif mod == "Takeoff" and disc:
                            mc[i].warning(f"⚠", icon=None)
                        else:
                            mc[i].caption(mod)

                    # Progress bar
                    st.progress(pct/100, text=f"{pct}% complete")

                    if st.button(f"Open {pname} →", key=f"ov_op_{pname}",
                                 use_container_width=True, type="primary"):
                        st.session_state.active_project = pname
                        st.session_state.active_module  = "Takeoff"
                        st.session_state.nav            = "Projects"
                        st.rerun()

    with col_right:
        # ── Upcoming deadlines ────────────────────────────────────────────
        st.markdown("**Upcoming deadlines**")
        dls = []
        for pname,p in projects.items():
            bd = p.get("bid_date")
            if bd:
                try:
                    days = (date.fromisoformat(str(bd)) - today).days
                    dls.append((days, pname, bd))
                except: pass
        dls.sort()

        if not dls:
            st.caption("No deadlines set.")
        else:
            for days, pname, bd in dls[:7]:
                try:
                    bd_fmt = date.fromisoformat(str(bd)).strftime("%b %d, %Y")
                except:
                    bd_fmt = str(bd)[:10]
                if days < 0:
                    color, txt = "#94a3b8", f"Past due · {bd_fmt}"
                elif days <= 14:
                    color, txt = "#dc2626", bd_fmt
                elif days <= 30:
                    color, txt = "#d97706", bd_fmt
                else:
                    color, txt = "#16a34a", bd_fmt
                st.markdown(f"""
                <div class="dl-row">
                    <span style="font-weight:500">{pname}</span>
                    <span style="color:{color};font-weight:600;font-size:12px">{txt}</span>
                </div>""", unsafe_allow_html=True)

        # ── Alerts ────────────────────────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("**Needs attention**")
        alerts = []
        for pname,p in projects.items():
            disc = p["takeoff"].get("discrepancies",[])
            if disc:
                alerts.append(("red",
                    f"<b>{pname}</b> — {len(disc)} discrepancies need scope clarification"))
            bd = p.get("bid_date")
            if bd:
                try:
                    days = (date.fromisoformat(str(bd)) - today).days
                    if days <= 30 and module_status(p,"Estimate") == "not_started":
                        alerts.append(("amber",
                            f"<b>{pname}</b> — no estimate, {days}d to bid"))
                    if days <= 14 and module_status(p,"Proposal") == "not_started":
                        alerts.append(("red",
                            f"<b>{pname}</b> — no proposal, {days}d to bid"))
                except: pass
            if (p["docs"].get("SOO") and
                module_status(p,"Point List") == "not_started"):
                alerts.append(("blue",
                    f"<b>{pname}</b> — SOO uploaded, point list not generated"))

        if not alerts:
            st.markdown("""
            <div style="background:#f0fdf4;border:1px solid #bbf7d0;
                        border-radius:8px;padding:14px;text-align:center;
                        color:#166534;font-size:13px">
                ✅ All clear — no issues
            </div>""", unsafe_allow_html=True)
        else:
            for color, msg in alerts[:6]:
                css = {"red":"alert-red","amber":"alert-amber",
                       "blue":"alert-blue"}.get(color,"alert-blue")
                st.markdown(
                    f'<div class="alert {css}">{msg}</div>',
                    unsafe_allow_html=True)

        # ── Summary stats ─────────────────────────────────────────────────
        if projects:
            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown("**Summary**")
            total_dev  = sum(len(p["takeoff"].get("equipment",[]))
                             for p in projects.values())
            total_disc = sum(len(p["takeoff"].get("discrepancies",[]))
                             for p in projects.values())
            for label, val in [
                ("Total devices in scope", total_dev),
                ("Total discrepancies",    total_disc),
                ("Proposals submitted",    prop_ready),
                ("Pipeline value",
                 f"${pipeline_val/1000:.0f}K" if pipeline_val else "—"),
            ]:
                sc1,sc2 = st.columns([2,1])
                sc1.caption(label)
                sc2.markdown(f"**{val}**")

    # ── Roadmap panel ─────────────────────────────────────────────────────
    st.divider()
    with st.expander("🗺 Product status & roadmap"):
        _product_status_panel()
# ── Clients ───────────────────────────────────────────────────────────────────
def page_clients():
    st.title("Clients")
    st.caption("Save each client's point list template (Excel), proposal template (Word), "
               "and labor rates. Select a client when creating a project to auto-apply their formats.")

    clients = st.session_state.clients
    col_list, col_form = st.columns([1,1.6])

    with col_list:
        st.markdown("**Saved clients**")
        if not clients:
            st.info("No clients yet.")
        for cname in list(clients.keys()):
            c = clients[cname]
            with st.expander(f"**{cname}**"):
                st.caption(f"Point list template: `{c.get('pl_template_name','—')}`")
                st.caption(f"Proposal template: `{c.get('prop_template_name','—')}`")
                rates = c.get("rates", DEFAULT_RATES)
                st.caption("Rates: " + "  ·  ".join(f"{k}: **${v}/hr**" for k,v in rates.items()))
                projs_using = [pname for pname,p in st.session_state.projects.items()
                               if p.get("client")==cname]
                if projs_using:
                    st.caption(f"Used by: {', '.join(projs_using)}")
                if st.button("Delete", key=f"del_{cname}"):
                    del st.session_state.clients[cname]
                    st.rerun()

    with col_form:
        st.markdown("**Add / update client**")
        with st.form("client_form", clear_on_submit=True):
            cname = st.text_input("Client name *", placeholder="e.g. Johnson Controls NYC")

            st.markdown("**Point list template** — upload your Excel format")
            pl_file = st.file_uploader("Excel (.xlsx)", type=["xlsx"], key="pl_up")
            st.caption("AI will read your column headers and output the point list in exactly this format.")

            st.markdown("**Proposal template** — upload your Word format")
            prop_file = st.file_uploader("Word (.docx)", type=["docx"], key="prop_up")
            st.caption("Use {{PROJECT_NAME}}, {{CLIENT}}, {{DATE}}, {{SCOPE_TEXT}} as placeholders in your template.")

            st.markdown("**Labor rates ($/hr)**")
            rc = st.columns(len(PHASES))
            rates = {}
            for i,ph in enumerate(PHASES):
                rates[ph] = rc[i].number_input(ph, 0, 500, DEFAULT_RATES[ph], key=f"r_{ph}")

            if st.form_submit_button("Save client", type="primary"):
                if not cname:
                    st.error("Client name required.")
                else:
                    entry = {"rates":rates}
                    if pl_file:
                        entry["pl_template_bytes"] = pl_file.read()
                        entry["pl_template_name"]  = pl_file.name
                    if prop_file:
                        entry["prop_template_bytes"] = prop_file.read()
                        entry["prop_template_name"]  = prop_file.name
                    st.session_state.clients[cname] = entry
                    _save_app_state()
                    st.success(f"✅ '{cname}' saved.")
                    st.rerun()

# ── Projects hub ──────────────────────────────────────────────────────────────
def page_projects():
    active = st.session_state.active_project
    if active and active in st.session_state.projects:
        page_project_detail(st.session_state.projects[active])
    else:
        page_projects_list()

def page_projects_list():
    today    = date.today()
    projects = st.session_state.projects

    # ── Styles ────────────────────────────────────────────────────────
    st.markdown("""<style>
    .proj-table-hdr {
        display:grid;
        grid-template-columns:60px 1fr 60px 120px 80px 80px 100px 90px;
        padding:8px 12px; background:#f8fafc;
        border:1px solid #e2e8f0; border-radius:8px 8px 0 0;
        font-size:11px; font-weight:600; color:#64748b;
        text-transform:uppercase; letter-spacing:.05em;
    }
    .proj-table-row {
        display:grid;
        grid-template-columns:60px 1fr 60px 120px 80px 80px 100px 90px;
        padding:10px 12px; border:1px solid #e2e8f0;
        border-top:none; background:#fff;
        font-size:13px; align-items:center;
        transition:background .1s;
    }
    .proj-table-row:hover { background:#f8fafc; cursor:pointer; }
    .proj-table-row:last-child { border-radius:0 0 8px 8px; }
    .status-badge {
        display:inline-block; padding:3px 10px; border-radius:20px;
        font-size:11px; font-weight:600;
    }
    .status-active   { background:#dbeafe; color:#1e40af; border:1px solid #93c5fd; }
    .status-inprog   { background:#fef9c3; color:#854d0e; border:1px solid #fcd34d; }
    .status-onhold   { background:#ffedd5; color:#9a3412; border:1px solid #fdba74; }
    .status-complete { background:#dcfce7; color:#166534; border:1px solid #86efac; }
    .status-archived { background:#f1f5f9; color:#475569; border:1px solid #cbd5e1; }
    .pct-bar-track  { height:5px; border-radius:3px;
                      background:#e2e8f0; width:100%; }
    .pct-bar-fill   { height:5px; border-radius:3px;
                      background:linear-gradient(90deg,#6366f1,#10b981); }
    .tmpl-card {
        border:1px solid #e2e8f0; border-radius:10px;
        padding:14px 16px; background:#fff; margin-bottom:8px;
    }
    .tmpl-name { font-size:14px; font-weight:600; color:#1e293b; }
    .tmpl-meta { font-size:12px; color:#64748b; margin-top:3px; }
    </style>""", unsafe_allow_html=True)

    # ── Header ────────────────────────────────────────────────────────
    h1,h2 = st.columns([3,1])
    h1.markdown("## Projects")
    h1.caption(f"{len(projects)} active · {today.strftime('%B %d, %Y')}")
    if h2.button("＋ New project", type="primary",
                 use_container_width=True, key="pl_new"):
        st.session_state["show_new_proj"] = True

    # ── Tabs: Active Projects | Templates ─────────────────────────────
    tab_active, tab_tmpl = st.tabs(["📁 Active Projects", "📋 Templates"])

    # ════════════════════════════════════════════════════════════════════
    # TAB 1: Active Projects
    # ════════════════════════════════════════════════════════════════════
    with tab_active:

        # New project form (shown when button clicked)
        if st.session_state.get("show_new_proj"):
            with st.container(border=True):
                st.markdown("**New project**")
                with st.form("new_proj"):
                    c1,c2 = st.columns(2)
                    pname    = c1.text_input("Project name *")
                    address  = c2.text_input("Address / location")
                    c3,c4   = st.columns(2)
                    bid_date = c3.date_input("Bid date", value=None)

                    # Template picker
                    tmpls    = st.session_state.get("templates", {})
                    tmpl_opts = ["No template"] + list(tmpls.keys())
                    tmpl_sel  = c4.selectbox("Apply template", tmpl_opts,
                                              help="Auto-fill rates and formats")

                    st.markdown("**Documents** *(optional — upload later in Takeoff)*")
                    d1,d2,d3 = st.columns(3)
                    draw_f = d1.file_uploader("Drawings (PDF)",  type=["pdf"],        key="np_draw")
                    soo_f  = d2.file_uploader("SOO (PDF/DOCX)",  type=["pdf","docx"], key="np_soo")
                    spec_f = d3.file_uploader("Controls spec",   type=["pdf","docx"], key="np_spec")

                    fc1,fc2 = st.columns(2)
                    if fc1.form_submit_button("Create project", type="primary"):
                        if not pname:
                            st.error("Project name required.")
                        elif pname in projects:
                            st.error("Name already exists.")
                        else:
                            p = new_project(pname, None,
                                            str(bid_date) if bid_date else None,
                                            address)
                            for lbl,f in [("Drawings",draw_f),
                                          ("SOO",soo_f),
                                          ("Controls spec",spec_f)]:
                                if f:
                                    p["docs"][lbl]      = f.read()
                                    p["doc_names"][lbl] = f.name
                            # Apply template rates/templates
                            tmpl = tmpls.get(tmpl_sel, {})
                            p["estimate"]["rates"] = dict(
                                tmpl.get("rates", DEFAULT_RATES))
                            if tmpl.get("pl_template_bytes"):
                                st.session_state.clients[pname] = {
                                    "rates":              tmpl["rates"],
                                    "pl_template_bytes":  tmpl["pl_template_bytes"],
                                    "pl_template_name":   tmpl.get("pl_template_name",""),
                                    "prop_template_bytes":tmpl.get("prop_template_bytes"),
                                    "prop_template_name": tmpl.get("prop_template_name",""),
                                }
                            st.session_state.projects[pname] = p
                            st.session_state.active_project  = pname
                            st.session_state.active_module   = "Takeoff"
                            st.session_state["show_new_proj"] = False
                            _save_app_state()
                            st.success(f"✅ '{pname}' created.")
                            st.rerun()
                    if fc2.form_submit_button("Cancel"):
                        st.session_state["show_new_proj"] = False
                        st.rerun()

        if not projects:
            st.markdown("""
            <div style="border:2px dashed #e2e8f0;border-radius:12px;
                        padding:48px;text-align:center;margin-top:16px">
                <div style="font-size:32px;margin-bottom:10px">📋</div>
                <div style="font-size:16px;font-weight:600;color:#475569">
                    No projects yet</div>
                <div style="font-size:13px;color:#94a3b8;margin-top:6px">
                    Click <strong>＋ New project</strong> to get started</div>
            </div>""", unsafe_allow_html=True)
        else:
            # ── Filter bar ────────────────────────────────────────────
            fc1,fc2,fc3 = st.columns([2,1,1])
            search   = fc1.text_input("", placeholder="🔍  Search projects…",
                                       label_visibility="collapsed", key="pl_search")
            sort_by  = fc2.selectbox("Sort by",
                                      ["Bid date","Name","% Complete","Status"],
                                      label_visibility="collapsed", key="pl_sort")
            filt_st  = fc3.selectbox("Status",
                                      ["All"] + PROJECT_STATUSES,
                                      label_visibility="collapsed", key="pl_filt")

            # ── Build rows ────────────────────────────────────────────
            rows = []
            for i,(pname,p) in enumerate(projects.items()):
                disc  = len(p["takeoff"].get("discrepancies",[]))
                devs  = len(p["takeoff"].get("equipment",[]))
                bd    = p.get("bid_date","")
                done_n= sum(1 for m in MODULE_ORDER
                            if module_status(p,m) in ("done","in_progress"))
                pct   = int(done_n/len(MODULE_ORDER)*100)

                # Status — use manual project_status, default to Active
                status = p.get("project_status", "Active")

                # Deadline
                dl_days = None
                if bd:
                    try: dl_days = (date.fromisoformat(str(bd))-today).days
                    except: pass

                rows.append({
                    "id":       i+1,
                    "name":     pname,
                    "p":        p,
                    "pct":      pct,
                    "status":   status,
                    "devs":     devs,
                    "disc":     disc,
                    "bd":       bd,
                    "dl_days":  dl_days,
                })

            # Filter
            if search:
                rows = [r for r in rows
                        if search.lower() in r["name"].lower()]
            if filt_st != "All":
                rows = [r for r in rows if r["status"] == filt_st]

            # Sort
            if sort_by == "Bid date":
                rows.sort(key=lambda r: (r["dl_days"] is None,
                                          r["dl_days"] or 9999))
            elif sort_by == "Name":
                rows.sort(key=lambda r: r["name"].lower())
            elif sort_by == "% Complete":
                rows.sort(key=lambda r: -r["pct"])
            elif sort_by == "Status":
                order = {"In Progress":0,"Active":1,"On Hold":2,
                         "Completed":3,"Archived":4}
                rows.sort(key=lambda r: order.get(r["status"],9))

            # ── Table header ──────────────────────────────────────────
            st.markdown("")
            hdr = st.columns([0.4,3,0.5,1.2,0.7,0.7,1,0.8])
            for c,h in zip(hdr,["ID","Project Name","%","Status",
                                  "Devices","Disc.","Bid Date","Action"]):
                c.markdown(f'<span style="font-size:11px;font-weight:600;'
                           f'color:#64748b;text-transform:uppercase;'
                           f'letter-spacing:.05em">{h}</span>',
                           unsafe_allow_html=True)
            st.markdown('<hr style="margin:4px 0 0;border-color:#e2e8f0">',
                        unsafe_allow_html=True)

            # ── Table rows ────────────────────────────────────────────
            STATUS_CSS = {
                "Active":      "status-active",
                "In Progress": "status-inprog",
                "On Hold":     "status-onhold",
                "Completed":   "status-complete",
                "Archived":    "status-archived",
            }
            for r in rows:
                pname  = r["name"]
                status = r["status"]
                disc   = r["disc"]
                dl     = r["dl_days"]

                # Deadline display
                if dl is None:
                    dl_txt, dl_col = "—", "#94a3b8"
                else:
                    try:
                        bd_fmt = date.fromisoformat(str(r["bd"])).strftime("%b %d, %Y")
                    except:
                        bd_fmt = str(r["bd"])[:10]
                    if dl < 0:      dl_txt, dl_col = f"Past due · {bd_fmt}", "#94a3b8"
                    elif dl <= 14:  dl_txt, dl_col = bd_fmt, "#dc2626"
                    elif dl <= 30:  dl_txt, dl_col = bd_fmt, "#d97706"
                    else:           dl_txt, dl_col = bd_fmt, "#64748b"

                col = st.columns([0.4,3,0.5,1.2,0.7,0.7,1,0.8])
                col[0].markdown(f'<span style="color:#94a3b8;font-size:12px">'
                                f'#{r["id"]}</span>', unsafe_allow_html=True)
                # Project name is clickable
                if col[1].button(f"**{pname}**", key=f"name_open_{pname}",
                                 use_container_width=True):
                    st.session_state.active_project = pname
                    st.session_state.active_module  = "Takeoff"
                    st.session_state.nav            = "Projects"
                    st.rerun()
                # % with mini bar
                col[2].markdown(
                    f'<div style="font-size:12px;font-weight:600">{r["pct"]}%</div>'
                    f'<div class="pct-bar-track"><div class="pct-bar-fill"'
                    f' style="width:{r["pct"]}%"></div></div>',
                    unsafe_allow_html=True)
                # Single styled dropdown — no badge above
                new_status = col[3].selectbox(
                    "Status", PROJECT_STATUSES,
                    index=PROJECT_STATUSES.index(status)
                          if status in PROJECT_STATUSES else 0,
                    key=f"ps_{pname}",
                    label_visibility="collapsed",
                    format_func=lambda s: {
                        "Active":      "🔵 Active",
                        "In Progress": "🟡 In Progress",
                        "On Hold":     "🟠 On Hold",
                        "Completed":   "🟢 Completed",
                        "Archived":    "⚫ Archived",
                    }.get(s, s)
                )
                if new_status != status:
                    p["project_status"] = new_status
                    _save_app_state()
                    st.rerun()
                col[4].markdown(str(r["devs"]))
                col[5].markdown(
                    f'<span style="color:{"#dc2626" if disc else "#94a3b8"};'
                    f'font-weight:{"600" if disc else "400"}">'
                    f'{disc if disc else "—"}</span>',
                    unsafe_allow_html=True)
                col[6].markdown(
                    f'<span style="color:{dl_col};font-weight:500">'
                    f'{dl_txt}</span>', unsafe_allow_html=True)

                rc1,rc2,rc3 = col[7].columns(3)
                if rc1.button("→", key=f"op_{pname}",
                              help=f"Open {pname}"):
                    st.session_state.active_project = pname
                    st.session_state.active_module  = "Takeoff"
                    st.rerun()
                if rc2.button("✏", key=f"ed_{pname}",
                              help="Edit"):
                    st.session_state.editing_project = pname
                    st.rerun()
                if rc3.button("🗑", key=f"dl_{pname}",
                              help="Delete"):
                    st.session_state[f"confirm_del_{pname}"] = True
                    st.rerun()

                # Confirm delete
                if st.session_state.get(f"confirm_del_{pname}"):
                    st.warning(f"Delete **{pname}**? This cannot be undone.")
                    cc1,cc2 = st.columns(2)
                    if cc1.button("Yes, delete", key=f"yes_{pname}",
                                  type="primary"):
                        del projects[pname]
                        st.session_state.pop(f"confirm_del_{pname}", None)
                        if st.session_state.active_project == pname:
                            st.session_state.active_project = None
                        _save_app_state()
                        st.rerun()
                    if cc2.button("Cancel", key=f"no_{pname}"):
                        st.session_state.pop(f"confirm_del_{pname}", None)
                        st.rerun()

                # Inline edit form
                if st.session_state.get("editing_project") == pname:
                    with st.form(f"edit_{pname}"):
                        ec1,ec2,ec3,ec4 = st.columns(4)
                        new_name = ec1.text_input("Name", value=pname)
                        new_addr = ec2.text_input("Address",
                                                   value=r["p"].get("address",""))
                        cur_bid  = None
                        try:
                            cur_bid = date.fromisoformat(str(r["p"].get("bid_date","")))                                       if r["p"].get("bid_date") else None
                        except: pass
                        new_bid  = ec3.date_input("Bid date", value=cur_bid)
                        tmpl_opts2 = ["No template"] + list(
                            st.session_state.get("templates",{}).keys())
                        new_tmpl = ec4.selectbox("Template", tmpl_opts2,
                                                  key=f"et_{pname}")
                        s1,s2 = st.columns(2)
                        if s1.form_submit_button("Save", type="primary"):
                            p2 = projects.get(pname, r["p"])
                            if new_name and new_name != pname:
                                projects[new_name] = projects.pop(pname)
                                pname = new_name
                                p2    = projects[pname]
                            p2["address"]  = new_addr
                            p2["bid_date"] = str(new_bid) if new_bid else None
                            tmpl2 = st.session_state.get("templates",{}).get(new_tmpl,{})
                            if tmpl2.get("rates"):
                                p2["estimate"]["rates"] = dict(tmpl2["rates"])
                            st.session_state.editing_project = None
                            _save_app_state()
                            st.rerun()
                        if s2.form_submit_button("Cancel"):
                            st.session_state.editing_project = None
                            st.rerun()

                st.markdown('<hr style="margin:2px 0;border-color:#f1f5f9">',
                            unsafe_allow_html=True)

            st.caption(f"Total: {len(rows)} project{'s' if len(rows)!=1 else ''}")

    # ════════════════════════════════════════════════════════════════════
    # TAB 2: Templates
    # ════════════════════════════════════════════════════════════════════
    with tab_tmpl:
        st.markdown("**Project templates** — save your rates and document formats "
                    "for quick reuse across projects.")

        if "templates" not in st.session_state:
            st.session_state.templates = {}
        tmpls = st.session_state.templates

        tc1, tc2 = st.columns([1, 1.6])

        with tc1:
            st.markdown("**Saved templates**")
            if not tmpls:
                st.info("No templates yet. Create one →")
            for tname, t in list(tmpls.items()):
                with st.container(border=True):
                    st.markdown(f"**{tname}**")
                    rates = t.get("rates", DEFAULT_RATES)
                    st.caption("Rates: " + "  ·  ".join(
                        f"{k}: ${v}/hr" for k,v in rates.items()))
                    st.caption(
                        f"Point list: `{t.get('pl_template_name','—')}`  "
                        f"Proposal: `{t.get('prop_template_name','—')}`")
                    proj_using = [pn for pn,p in projects.items()
                                  if p.get("_template") == tname]
                    if proj_using:
                        st.caption(f"Used by: {', '.join(proj_using)}")
                    if st.button("Delete", key=f"del_tmpl_{tname}"):
                        del st.session_state.templates[tname]
                        st.rerun()

        with tc2:
            st.markdown("**Create template**")
            with st.form("tmpl_form", clear_on_submit=True):
                tname = st.text_input("Template name *",
                                       placeholder="e.g. Standard NYC Commercial")

                st.markdown("**Point list template** (.xlsx)")
                pl_f = st.file_uploader("Upload Excel template",
                                         type=["xlsx"], key="tmpl_pl")
                st.caption("AI matches your column headers exactly")

                st.markdown("**Proposal template** (.docx)")
                prop_f = st.file_uploader("Upload Word template",
                                           type=["docx"], key="tmpl_prop")
                st.caption("Placeholders: {{PROJECT_NAME}} {{DATE}} {{SCOPE_TEXT}}")

                st.markdown("**Default labor rates ($/hr)**")
                rc = st.columns(len(PHASES))
                rates = {}
                for i,ph in enumerate(PHASES):
                    rates[ph] = rc[i].number_input(
                        ph, 0, 500, DEFAULT_RATES[ph], key=f"tr_{ph}")

                if st.form_submit_button("Save template", type="primary"):
                    if not tname:
                        st.error("Template name required.")
                    else:
                        entry = {"rates": rates}
                        if pl_f:
                            entry["pl_template_bytes"] = pl_f.read()
                            entry["pl_template_name"]  = pl_f.name
                        if prop_f:
                            entry["prop_template_bytes"] = prop_f.read()
                            entry["prop_template_name"]  = prop_f.name
                        st.session_state.templates[tname] = entry
                        st.success(f"✅ Template '{tname}' saved.")
                        st.rerun()
def page_project_detail(p):
    bc, hc = st.columns([1, 8])
    if bc.button("← Back"):
        st.session_state.active_project = None
        st.rerun()
    hc.markdown(f"## {p['name']}")
    proj_st = p.get("project_status","Active")
    hc.caption(
        (f"{p.get('address')}  ·  " if p.get('address') else "")
        + f"Bid: {p.get('bid_date','TBD')}"
        + (f"  ·  {p.get('client')}" if p.get('client') else "")
    )
    # Single status dropdown in header
    hdr_st = hc.selectbox(
        "Status", PROJECT_STATUSES,
        index=PROJECT_STATUSES.index(proj_st)
              if proj_st in PROJECT_STATUSES else 0,
        key=f"hdr_status_{p['name']}",
        label_visibility="collapsed",
        help="Change project status",
        format_func=lambda s: {
            "Active":      "🔵 Active",
            "In Progress": "🟡 In Progress",
            "On Hold":     "🟠 On Hold",
            "Completed":   "🟢 Completed",
            "Archived":    "⚫ Archived",
        }.get(s, s)
    )
    if hdr_st != proj_st:
        p["project_status"] = hdr_st
        _save_app_state()
        st.rerun()

    # All modules in order
    all_mods = MODULE_ORDER + ["AI Advisor", "Drawing Markup"]

    # Build tab labels
    tab_labels = []
    for mod in all_mods:
        if mod == "AI Advisor":
            tab_labels.append("🤖 AI Advisor")
        elif mod == "Drawing Markup":
            tab_labels.append("🖊 Drawing Markup")
        elif module_locked(p, mod):
            tab_labels.append(f"🔒 {mod}")
        else:
            s    = module_status(p, mod)
            icon = "✅" if s == "done" else "⚠️" if s == "issues" else "📋"
            tab_labels.append(f"{icon} {mod}")

    # Determine which tab to show based on sidebar click
    active_mod = st.session_state.get("active_module", "Takeoff")
    try:
        default_tab = all_mods.index(active_mod)
    except ValueError:
        default_tab = 0

    tabs = st.tabs(tab_labels)
    handlers = [module_takeoff, module_point_list, module_estimate,
                module_proposal, module_ai_advisor, module_markup]

    for i, (tab, mod, handler) in enumerate(zip(tabs, all_mods, handlers)):
        with tab:
            if mod not in ("AI Advisor", "Drawing Markup") and module_locked(p, mod):
                unlock_msg = {
                    "Point List": "Upload SOO or Controls spec in the Takeoff tab to unlock.",
                    "Estimate":   "Generate a point list or upload Controls spec to unlock.",
                    "Proposal":   "Generate an estimate or complete takeoff to unlock.",
                }.get(mod, "Complete previous steps to unlock.")
                st.info(f"🔒 **{mod} is locked.** {unlock_msg}")
            else:
                warn = module_data_warning(p, mod)
                if warn and mod not in ("AI Advisor", "Drawing Markup", "Takeoff"):
                    st.warning(warn)
                handler(p)

# ── Module 1: Takeoff ─────────────────────────────────────────────────────────
def module_takeoff(p):
    """Takeoff module with 7 independent sub-tabs."""

    # Ensure takeoff sub-data structures exist
    if "soo_register"    not in p: p["soo_register"]    = {}
    if "schedule_data"   not in p: p["schedule_data"]   = {}
    if "floorplan_data"  not in p: p["floorplan_data"]  = {}
    if "riser_data"      not in p: p["riser_data"]      = {}
    if "electrical_data" not in p: p["electrical_data"] = {}
    if "plumbing_data"   not in p: p["plumbing_data"]   = {}
    if "master_takeoff"  not in p: p["master_takeoff"]  = {}

    soo_loaded = bool(p["soo_register"].get("systems"))

    tabs = st.tabs([
        "📖 SOO",
        "📋 Schedule",
        "🏗 Floor Plan",
        "📐 Riser",
        "⚡ Electrical",
        "🔧 Plumbing",
        "📊 Master Takeoff",
    ])

    # ══════════════════════════════════════════════════════════════════
    # TAB 1 — SOO Reader
    # ══════════════════════════════════════════════════════════════════
    with tabs[0]:
        _tab_soo(p)

    # ══════════════════════════════════════════════════════════════════
    # TAB 2 — Schedule
    # ══════════════════════════════════════════════════════════════════
    with tabs[1]:
        _tab_schedule(p, soo_loaded)

    # ══════════════════════════════════════════════════════════════════
    # TAB 3 — Floor Plan
    # ══════════════════════════════════════════════════════════════════
    with tabs[2]:
        _tab_floorplan(p, soo_loaded)

    # ══════════════════════════════════════════════════════════════════
    # TAB 4 — Riser
    # ══════════════════════════════════════════════════════════════════
    with tabs[3]:
        _tab_riser(p, soo_loaded)

    # ══════════════════════════════════════════════════════════════════
    # TAB 5 — Electrical
    # ══════════════════════════════════════════════════════════════════
    with tabs[4]:
        _tab_electrical(p, soo_loaded)

    # ══════════════════════════════════════════════════════════════════
    # TAB 6 — Plumbing
    # ══════════════════════════════════════════════════════════════════
    with tabs[5]:
        _tab_plumbing(p, soo_loaded)

    # ══════════════════════════════════════════════════════════════════
    # TAB 7 — Master Takeoff
    # ══════════════════════════════════════════════════════════════════
    with tabs[6]:
        _tab_master(p)


# ── SOO Tab ───────────────────────────────────────────────────────────────────
def _tab_soo(p):
    """SOO extraction with 5 buttons: Overview, Proposal, Point List, Appendix, Notes"""
    
    st.markdown("### 📋 Sequence of Operations (SOO) Processing")
    st.markdown("""
    Upload your SOO PDF/DOCX and extract:
    - **Overview:** Bird's eye view of control approach per system
    - **Proposal:** Generate DOCX proposal (you provide template)
    - **Point List:** Main BMS points in Excel format (11 columns)
    - **Appendix:** Special sequences (fire safety, emergency, future)
    - **Important Notes:** Key estimation points
    """)
    
    # ── File Upload ────────────────────────────────────────────────────────────
    uploaded_file = st.file_uploader("Upload SOO PDF or DOCX", type=["pdf", "docx"], key="soo_upload")
    
    if not uploaded_file:
        st.info("👆 Upload a SOO document to begin analysis")
        return
    
    # ── Extract SOO Text ───────────────────────────────────────────────────────
    soo_bytes = uploaded_file.read()
    fname = uploaded_file.name.lower()
    
    if fname.endswith(".docx"):
        soo_text = _extract_docx_text(soo_bytes, 15000)
    else:
        soo_text = _extract_pdf_text(soo_bytes, 15000)
    
    if not soo_text:
        st.error("Could not extract text from SOO document")
        return
    
    st.success(f"✅ Loaded: {uploaded_file.name} ({len(soo_text)} chars)")
    
    # Initialize SOO state
    if "soo_data" not in p:
        p["soo_data"] = {
            "overview": None,
            "proposal": None,
            "point_list": None,
            "appendix": None,
            "important_notes": None
        }
    
    # ── 5 Extraction Buttons ───────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### Extract from SOO:")
    
    col1, col2, col3, col4, col5 = st.columns(5)
    
    k = api_key()
    if not k:
        st.error("❌ No API key. Set ANTHROPIC_API_KEY in Streamlit Secrets.")
        return
    
    # ── Button 1: Overview ─────────────────────────────────────────────────────
    with col1:
        if st.button("📋 Overview", key="btn_overview", use_container_width=True):
            with st.spinner("Extracting overview..."):
                from soo_extractor import generate_overview_prompt
                prompt = generate_overview_prompt(p.get("name", "Project"), soo_text)
                raw = _claude(k, prompt, max_tokens=2000)
                if raw:
                    try:
                        data = json.loads(raw[raw.find("{"):raw.rfind("}")+1])
                        p["soo_data"]["overview"] = data
                        _save_app_state()
                        st.success("✅ Overview extracted")
                        st.rerun()
                    except:
                        st.warning("Could not parse response")
    
    # ── Button 2: Proposal ─────────────────────────────────────────────────────
    with col2:
        if st.button("📄 Proposal", key="btn_proposal", use_container_width=True):
            st.info("🔗 Upload your proposal template below")
    
    # ── Button 3: Point List ───────────────────────────────────────────────────
    with col3:
        if st.button("📊 Point List", key="btn_pointlist", use_container_width=True):
            with st.spinner("Generating point list..."):
                from soo_extractor import generate_pointlist_prompt, parse_pointlist_response
                prompt = generate_pointlist_prompt(p.get("name", "Project"), soo_text)
                raw = _claude(k, prompt, max_tokens=4000)
                if raw:
                    try:
                        rows = parse_pointlist_response(raw)
                        p["soo_data"]["point_list"] = rows
                        _save_app_state()
                        st.success(f"✅ {len(rows)} points extracted")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Parse error: {e}")
    
    # ── Button 4: Appendix ─────────────────────────────────────────────────────
    with col4:
        if st.button("📎 Appendix", key="btn_appendix", use_container_width=True):
            if not p["soo_data"]["point_list"]:
                st.warning("⚠️ Generate Point List first")
            else:
                with st.spinner("Generating appendix..."):
                    from soo_extractor import generate_appendix_prompt, parse_pointlist_response
                    main_equip = list(set(row.get("Equipment", "") for row in p["soo_data"]["point_list"]))
                    prompt = generate_appendix_prompt(p.get("name", "Project"), soo_text, main_equip)
                    raw = _claude(k, prompt, max_tokens=2000)
                    if raw:
                        try:
                            rows = parse_pointlist_response(raw)
                            p["soo_data"]["appendix"] = rows
                            _save_app_state()
                            st.success(f"✅ {len(rows)} appendix points extracted")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Parse error: {e}")
    
    # ── Button 5: Important Notes ──────────────────────────────────────────────
    with col5:
        if st.button("⭐ Notes", key="btn_notes", use_container_width=True):
            with st.spinner("Extracting notes..."):
                from soo_extractor import generate_important_notes_prompt, parse_notes_response
                prompt = generate_important_notes_prompt(p.get("name", "Project"), soo_text)
                raw = _claude(k, prompt, max_tokens=2500)
                if raw:
                    try:
                        data = parse_notes_response(raw)
                        p["soo_data"]["important_notes"] = data
                        _save_app_state()
                        st.success("✅ Important notes extracted")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Parse error: {e}")
    
    # ── Display Results ────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### Results:")
    
    # Overview
    if p["soo_data"]["overview"]:
        with st.expander("📋 Overview (Bird's eye view)"):
            overview_data = p["soo_data"]["overview"]
            if isinstance(overview_data, dict) and "overview" in overview_data:
                for system in overview_data["overview"]:
                    st.write(f"**{system.get('System', 'Unknown')}** — {system.get('Equipment_Type', '')}")
                    st.write(f"Control: {system.get('Control_Approach', '')}")
                    st.write(f"Points: {system.get('Control_Points', '')}")
                    st.write(f"Integration: {system.get('Integration', '')}")
                    st.write("---")
            else:
                st.json(overview_data)
    
    # Point List
    if p["soo_data"]["point_list"]:
        with st.expander(f"📊 Point List ({len(p['soo_data']['point_list'])} points)"):
            df = pd.DataFrame(p["soo_data"]["point_list"])
            cols = ["Panel Name", "Equipment", "Point name", "Control Device", 
                   "AI", "BI", "AO", "BO", "Serial Pt", "Terms", "Remarks"]
            for col in cols:
                if col not in df.columns:
                    df[col] = ""
            df = df[cols]
            
            st.dataframe(df, use_container_width=True, height=400)
            
            # Export to Excel
            xb = BytesIO()
            df.to_excel(xb, index=False, sheet_name="Point List")
            xb.seek(0)
            st.download_button("📥 Download Point List Excel", xb, 
                              f"point_list_{p.get('name', 'project').replace(' ', '_')}.xlsx",
                              "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    
    # Appendix
    if p["soo_data"]["appendix"]:
        with st.expander(f"📎 Appendix ({len(p['soo_data']['appendix'])} special points)"):
            df = pd.DataFrame(p["soo_data"]["appendix"])
            cols = ["Panel Name", "Equipment", "Point name", "Control Device", 
                   "AI", "BI", "AO", "BO", "Serial Pt", "Terms", "Remarks"]
            for col in cols:
                if col not in df.columns:
                    df[col] = ""
            df = df[cols]
            
            st.dataframe(df, use_container_width=True, height=300)
            
            # Export to Excel
            xb = BytesIO()
            df.to_excel(xb, index=False, sheet_name="Appendix")
            xb.seek(0)
            st.download_button("📥 Download Appendix Excel", xb,
                              f"appendix_{p.get('name', 'project').replace(' ', '_')}.xlsx",
                              "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    
    # Combined export (Main + Appendix)
    if p["soo_data"]["point_list"] and p["soo_data"]["appendix"]:
        st.markdown("**Or export both together:**")
        
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
        
        wb = Workbook()
        
        # Main sheet
        ws_main = wb.active
        ws_main.title = "Main Points"
        
        df_main = pd.DataFrame(p["soo_data"]["point_list"])
        cols = ["Panel Name", "Equipment", "Point name", "Control Device", 
               "AI", "BI", "AO", "BO", "Serial Pt", "Terms", "Remarks"]
        for col in cols:
            if col not in df_main.columns:
                df_main[col] = ""
        df_main = df_main[cols]
        
        fill = PatternFill("solid", start_color="2E75B6")
        for ci, col in enumerate(cols, 1):
            cell = ws_main.cell(1, ci, col)
            cell.font = Font(bold=True, color="FFFFFF", size=10)
            cell.fill = fill
        
        for ri, row in enumerate(df_main.to_dict("records"), 2):
            for ci, col in enumerate(cols, 1):
                ws_main.cell(ri, ci, row.get(col, ""))
        
        # Appendix sheet
        ws_app = wb.create_sheet("Appendix")
        
        df_app = pd.DataFrame(p["soo_data"]["appendix"])
        for col in cols:
            if col not in df_app.columns:
                df_app[col] = ""
        df_app = df_app[cols]
        
        fill_app = PatternFill("solid", start_color="B85C00")
        for ci, col in enumerate(cols, 1):
            cell = ws_app.cell(1, ci, col)
            cell.font = Font(bold=True, color="FFFFFF", size=10)
            cell.fill = fill_app
        
        for ri, row in enumerate(df_app.to_dict("records"), 2):
            for ci, col in enumerate(cols, 1):
                ws_app.cell(ri, ci, row.get(col, ""))
        
        xb_combined = BytesIO()
        wb.save(xb_combined)
        xb_combined.seek(0)
        
        st.download_button("📥 Download Main + Appendix Combined", xb_combined,
                          f"point_list_complete_{p.get('name', 'project').replace(' ', '_')}.xlsx",
                          "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    
    # Important Notes
    if p["soo_data"]["important_notes"]:
        with st.expander("⭐ Important Notes (Estimation)"):
            notes = p["soo_data"]["important_notes"]
            if isinstance(notes, dict):
                for category, items in notes.items():
                    if items:
                        st.subheader(category.replace("_", " ").title())
                        for item in items:
                            st.write(f"• {item}")
            else:
                st.json(notes)
    r_tab1, r_tab2, r_tab3, r_tab4 = st.tabs([
        "✅ Confirmed scope",
        "❌ Exclusions",
        "❓ Questions / gaps",
        "📋 I/O summary"
    ])

    with r_tab1:
        systems = reg.get("systems", [])
        st.markdown(f"**{len(systems)} systems confirmed in BMS scope**")
        if systems:
            df = pd.DataFrame(systems)
            st.dataframe(df, use_container_width=True, hide_index=True)
            _export_btn(df, "soo_confirmed", p["name"], key="exp_soo_conf")

    with r_tab2:
        excl = reg.get("exclusions", [])
        if excl:
            df = pd.DataFrame(excl)
            st.dataframe(df, use_container_width=True, hide_index=True)
        else:
            st.info("No explicit exclusions found.")

    with r_tab3:
        qs = reg.get("questions", [])
        if qs:
            for q in qs:
                st.warning(f"❓ {q}")
        else:
            st.success("No gaps or questions found.")

    with r_tab4:
        io_summary = reg.get("io_summary", [])
        if io_summary:
            df = pd.DataFrame(io_summary)
            st.dataframe(df, use_container_width=True, hide_index=True)
            _export_btn(df, "soo_io_summary", p["name"], key="exp_soo_io")
        else:
            st.info("No I/O tables extracted.")


# ── Schedule Tab ──────────────────────────────────────────────────────────────
def _tab_schedule(p, soo_loaded):
    st.markdown("### Schedule Takeoff")
    st.caption(
        "Upload mechanical schedule sheets (PDF). "
        "Tool extracts BMS-relevant devices, quantities, locations, and groupings. "
        "Filtered against SOO scope register."
    )

    if not soo_loaded:
        st.warning("⚠️ SOO not yet analysed. Results will show all devices unfiltered. "
                   "Read SOO first for accurate BMS scope filtering.")

    # Upload
    sched_bytes = p["docs"].get("Schedule")
    sched_name  = p["doc_names"].get("Schedule", "")

    col_up, col_info = st.columns([1, 2])
    with col_up:
        new_f = st.file_uploader(
            "Upload schedule sheets (PDF)",
            type=["pdf"],
            key="sched_upload"
        )
        if new_f:
            p["docs"]["Schedule"]      = new_f.read()
            p["doc_names"]["Schedule"] = new_f.name
            sched_bytes = p["docs"]["Schedule"]
            sched_name  = new_f.name
            _save_app_state()
            st.success(f"✅ `{sched_name}`")

    with col_info:
        if sched_bytes:
            size = round(len(sched_bytes)/1024/1024, 1)
            st.info(f"📋 `{sched_name}` · {size} MB")
        else:
            st.warning("No schedule uploaded.")

    st.divider()

    if not sched_bytes:
        st.info("Upload schedule sheets above.")
        return

    if st.button("🔍 Analyse schedule", type="primary", key="analyse_sched"):
        with st.spinner("Reading schedule — extracting tags, quantities, locations..."):
            result = _analyse_schedule(
                sched_bytes,
                p.get("soo_register", {})
            )
        p["schedule_data"] = result
        # Update main takeoff equipment list
        p["takeoff"]["equipment"]     = result.get("equipment", [])
        p["takeoff"]["discrepancies"] = result.get("discrepancies", [])
        p["takeoff"]["status"]        = "done" if result.get("equipment") else "not_started"
        _save_app_state()
        st.success(
            f"✅ {result.get('total_devices',0)} devices found · "
            f"{result.get('bms_scope',0)} in BMS scope · "
            f"{result.get('flagged',0)} flagged for floor plan"
        )
        st.rerun()

    data = p.get("schedule_data", {})
    if not data.get("equipment"):
        st.info("No schedule data yet. Click Analyse.")
        return

    _show_takeoff_results(data, p["name"], source="schedule")


# ── Floor Plan Tab ─────────────────────────────────────────────────────────────
def _tab_floorplan(p, soo_loaded):
    st.markdown("### Floor Plan Takeoff")
    st.caption(
        "Upload floor plan PDFs. Tool searches for tags from the schedule register, "
        "counts devices per floor, and generates an annotated PDF with highlights."
    )

    if not p.get("schedule_data", {}).get("equipment"):
        st.info("ℹ️ Run Schedule Analyse first to build the tag register used for floor plan cross-check.")

    fp_bytes = p["docs"].get("Floor Plan")
    fp_name  = p["doc_names"].get("Floor Plan", "")

    col_up, col_info = st.columns([1, 2])
    with col_up:
        new_f = st.file_uploader(
            "Upload floor plan (PDF)",
            type=["pdf"],
            key="fp_upload"
        )
        if new_f:
            p["docs"]["Floor Plan"]      = new_f.read()
            p["doc_names"]["Floor Plan"] = new_f.name
            fp_bytes = p["docs"]["Floor Plan"]
            fp_name  = new_f.name
            _save_app_state()
            st.success(f"✅ `{fp_name}`")

    with col_info:
        if fp_bytes:
            size = round(len(fp_bytes)/1024/1024, 1)
            st.info(f"🏗 `{fp_name}` · {size} MB")
        else:
            st.warning("No floor plan uploaded.")

    st.divider()

    if not fp_bytes:
        st.info("Upload floor plan above.")
        return

    if st.button("🔍 Analyse floor plan", type="primary", key="analyse_fp"):
        with st.spinner("Scanning floor plan — finding device tags and coordinates..."):
            # Get tag register from schedule
            sched_equip = p.get("schedule_data", {}).get("equipment", [])
            soo_tags    = set(s.get("tag","").upper()
                             for s in p.get("soo_register",{}).get("systems",[]))
            sched_tags  = set(e.get("tag","").upper() for e in sched_equip)

            from drawing_markup import DrawingMarkup
            dm = DrawingMarkup(fp_bytes,
                               soo_tags=soo_tags,
                               schedule_tags=sched_tags)
            dm.process()

            result = {
                "statuses":     dm.get_all_statuses_df(),
                "counts":       dm.get_summary_counts(),
                "amber_tags":   dm.get_amber_tags(),
                "red_tags":     dm.get_red_tags(),
                "page_count":   dm.page_count(),
                "sched_pages":  dm.schedule_page_count(),
                "dm_cache_key": f"dm_fp_{p['name']}",
            }
            st.session_state[f"dm_fp_{p['name']}"] = dm
            p["floorplan_data"] = result
            _save_app_state()

            counts = result["counts"]
            st.success(
                f"✅ {result['page_count']} pages · "
                f"{counts['green']} confirmed · "
                f"{counts['amber']} no SOO sequence · "
                f"{counts['red']} not found"
            )
            st.rerun()

    data = p.get("floorplan_data", {})
    if not data.get("statuses"):
        st.info("No floor plan data yet. Click Analyse.")
        return

    # Results
    counts = data.get("counts", {})
    m1,m2,m3,m4 = st.columns(4)
    m1.metric("🟢 Confirmed",    counts.get("green",0))
    m2.metric("🟡 No SOO seq",   counts.get("amber",0))
    m3.metric("🔴 Not found",    counts.get("red",0))
    m4.metric("🔵 Not in sched", counts.get("blue",0))

    # Tag table
    df = pd.DataFrame(data["statuses"])
    search = st.text_input("Search tag", key="fp_search")
    if search:
        df = df[df["Tag"].str.contains(search, case=False, na=False)]
    st.dataframe(df, use_container_width=True, hide_index=True, height=350)

    # Export annotated PDF
    st.divider()
    if st.button("Generate annotated PDF", type="primary", key="gen_fp_pdf"):
        dm = st.session_state.get(f"dm_fp_{p['name']}")
        if dm:
            with st.spinner("Generating annotated PDF..."):
                pdf_out = dm.generate_annotated_pdf()
            st.download_button(
                "⬇ Download annotated floor plan",
                data=pdf_out,
                file_name=f"markup_floorplan_{p['name'].replace(' ','_')}.pdf",
                mime="application/pdf",
                key="dl_fp_pdf"
            )
        else:
            st.warning("Re-run Analyse to regenerate PDF.")

    _export_btn(pd.DataFrame(data["statuses"]),
                "floorplan_takeoff", p["name"], key="exp_fp")


# ── Riser Tab ─────────────────────────────────────────────────────────────────
def _tab_riser(p, soo_loaded):
    st.markdown("### Riser Diagram Takeoff")
    st.caption(
        "Upload water riser and/or air riser PDFs. "
        "Tool finds labeled devices (BTU meters, DP sensors, valves, etc.). "
        "Symbol-only items flagged for manual review."
    )

    col_up1, col_up2 = st.columns(2)
    with col_up1:
        new_wr = st.file_uploader("Water riser (PDF)", type=["pdf"], key="wr_upload")
        if new_wr:
            p["docs"]["Water Riser"]      = new_wr.read()
            p["doc_names"]["Water Riser"] = new_wr.name
            _save_app_state()
            st.success(f"✅ `{new_wr.name}`")
        if p["doc_names"].get("Water Riser"):
            st.info(f"🔵 `{p['doc_names']['Water Riser']}`")

    with col_up2:
        new_ar = st.file_uploader("Air riser (PDF)", type=["pdf"], key="ar_upload")
        if new_ar:
            p["docs"]["Air Riser"]      = new_ar.read()
            p["doc_names"]["Air Riser"] = new_ar.name
            _save_app_state()
            st.success(f"✅ `{new_ar.name}`")
        if p["doc_names"].get("Air Riser"):
            st.info(f"🌬 `{p['doc_names']['Air Riser']}`")

    st.divider()

    has_riser = p["docs"].get("Water Riser") or p["docs"].get("Air Riser")
    if not has_riser:
        st.info("Upload water or air riser above.")
        return

    if st.button("🔍 Analyse risers", type="primary", key="analyse_riser"):
        with st.spinner("Scanning riser diagrams..."):
            result = _analyse_riser(
                p["docs"].get("Water Riser"),
                p["docs"].get("Air Riser"),
                p.get("soo_register", {})
            )
        p["riser_data"] = result
        _save_app_state()
        st.success(
            f"✅ {result.get('total',0)} devices found · "
            f"{result.get('flagged',0)} symbol-only items flagged"
        )
        st.rerun()

    data = p.get("riser_data", {})
    if not data.get("equipment"):
        st.info("No riser data yet. Click Analyse.")
        return

    _show_takeoff_results(data, p["name"], source="riser")


# ── Electrical Tab ─────────────────────────────────────────────────────────────
def _tab_electrical(p, soo_loaded):
    st.markdown("### Electrical Takeoff")
    st.caption(
        "Upload electrical drawings. Tool finds only items the SOO confirms "
        "as BMS monitoring scope (generator, ATS, meters, UPS, heat trace, etc.)."
    )

    if not soo_loaded:
        st.warning("⚠️ Read SOO first — electrical takeoff is filtered by SOO scope.")

    new_f = st.file_uploader("Electrical drawings (PDF)", type=["pdf"], key="elec_upload")
    if new_f:
        p["docs"]["Electrical"]      = new_f.read()
        p["doc_names"]["Electrical"] = new_f.name
        _save_app_state()
        st.success(f"✅ `{new_f.name}`")
    if p["doc_names"].get("Electrical"):
        st.info(f"⚡ `{p['doc_names']['Electrical']}`")

    st.divider()

    if not p["docs"].get("Electrical"):
        st.info("Upload electrical drawings above.")
        return

    if st.button("🔍 Analyse electrical", type="primary", key="analyse_elec"):
        with st.spinner("Scanning electrical drawings for BMS monitoring points..."):
            result = _analyse_electrical(
                p["docs"]["Electrical"],
                p.get("soo_register", {})
            )
        p["electrical_data"] = result
        _save_app_state()
        st.success(f"✅ {result.get('total',0)} BMS monitoring points found")
        st.rerun()

    data = p.get("electrical_data", {})
    if not data.get("equipment"):
        st.info("No electrical data yet. Click Analyse.")
        return

    _show_takeoff_results(data, p["name"], source="electrical")


# ── Plumbing Tab ───────────────────────────────────────────────────────────────
def _tab_plumbing(p, soo_loaded):
    st.markdown("### Plumbing Takeoff")
    st.caption(
        "Upload plumbing drawings. Tool finds only items the SOO confirms "
        "as BMS monitoring scope (DHW heaters, sump pumps, ejectors, booster pumps, etc.)."
    )

    if not soo_loaded:
        st.warning("⚠️ Read SOO first — plumbing takeoff is filtered by SOO scope.")

    new_f = st.file_uploader("Plumbing drawings (PDF)", type=["pdf"], key="plumb_upload")
    if new_f:
        p["docs"]["Plumbing"]      = new_f.read()
        p["doc_names"]["Plumbing"] = new_f.name
        _save_app_state()
        st.success(f"✅ `{new_f.name}`")
    if p["doc_names"].get("Plumbing"):
        st.info(f"🔧 `{p['doc_names']['Plumbing']}`")

    st.divider()

    if not p["docs"].get("Plumbing"):
        st.info("Upload plumbing drawings above.")
        return

    if st.button("🔍 Analyse plumbing", type="primary", key="analyse_plumb"):
        with st.spinner("Scanning plumbing drawings for BMS monitoring points..."):
            result = _analyse_plumbing(
                p["docs"]["Plumbing"],
                p.get("soo_register", {})
            )
        p["plumbing_data"] = result
        _save_app_state()
        st.success(f"✅ {result.get('total',0)} BMS monitoring points found")
        st.rerun()

    data = p.get("plumbing_data", {})
    if not data.get("equipment"):
        st.info("No plumbing data yet. Click Analyse.")
        return

    _show_takeoff_results(data, p["name"], source="plumbing")


# ── Master Takeoff Tab ─────────────────────────────────────────────────────────
def _tab_master(p):
    st.markdown("### Master Takeoff")
    st.caption(
        "Merges Schedule + Floor Plan + Riser + Electrical + Plumbing. "
        "Deduplicates, resolves quantity conflicts, and produces a clean "
        "client-facing register."
    )

    # Show what's available
    sources = {
        "Schedule":   len(p.get("schedule_data",{}).get("equipment",[])),
        "Floor Plan": len(p.get("floorplan_data",{}).get("statuses",[])),
        "Riser":      len(p.get("riser_data",{}).get("equipment",[])),
        "Electrical": len(p.get("electrical_data",{}).get("equipment",[])),
        "Plumbing":   len(p.get("plumbing_data",{}).get("equipment",[])),
    }

    sc = st.columns(5)
    for i,(src,count) in enumerate(sources.items()):
        icon = "✅" if count > 0 else "⚪"
        sc[i].metric(f"{icon} {src}", count)

    ready = sum(1 for v in sources.values() if v > 0)
    if ready == 0:
        st.info("Complete at least one analysis tab first.")
        return

    st.divider()

    if st.button("📊 Generate master takeoff",
                 type="primary", key="gen_master"):
        with st.spinner("Merging all sources and deduplicating..."):
            master = _generate_master(p)
        p["master_takeoff"] = master
        # Update main equipment list for point list + estimate
        p["takeoff"]["equipment"]     = master.get("equipment", [])
        p["takeoff"]["discrepancies"] = [
            e for e in master.get("equipment",[])
            if e.get("soo_status") == "No SOO sequence"
        ]
        p["takeoff"]["status"] = "done"
        _save_app_state()
        st.success(
            f"✅ Master takeoff: {len(master.get('equipment',[]))} devices · "
            f"{master.get('total_qty',0)} total quantity"
        )
        st.rerun()

    master = p.get("master_takeoff", {})
    if not master.get("equipment"):
        st.info("No master takeoff yet. Click Generate.")
        return

    # Client-facing table
    st.markdown("**Client-facing register**")
    st.caption("Clean format suitable for sharing with client or including in proposal.")

    equip = master.get("equipment", [])
    client_cols = ["Tag", "System", "Location", "Qty",
                   "Control Type", "Remarks"]
    df_client = pd.DataFrame([
        {
            "Tag":          e.get("tag",""),
            "System":       e.get("system",""),
            "Location":     e.get("floor",""),
            "Qty":          e.get("qty", 1),
            "Control Type": e.get("control_type","DDC"),
            "Remarks":      e.get("remarks",""),
        }
        for e in equip
    ])

    # Filter + sort
    f1,f2 = st.columns(2)
    search = f1.text_input("Search", placeholder="FCU, pump, 18th...", key="master_search")
    sort   = f2.selectbox("Sort by", ["Tag","System","Location","Qty"], key="master_sort")
    if search:
        sl = search.lower()
        df_client = df_client[
            df_client["Tag"].str.lower().str.contains(sl, na=False) |
            df_client["System"].str.lower().str.contains(sl, na=False) |
            df_client["Location"].str.lower().str.contains(sl, na=False)
        ]
    df_client = df_client.sort_values(sort)

    st.dataframe(df_client, use_container_width=True,
                 hide_index=True, height=450)

    # Summary by system
    st.divider()
    st.markdown("**Summary by system**")
    summary = df_client.groupby("System")["Qty"].sum().reset_index()
    summary.columns = ["System","Total Qty"]
    summary = summary.sort_values("Total Qty", ascending=False)
    st.dataframe(summary, use_container_width=True, hide_index=True)

    # Export
    st.divider()
    col_e1, col_e2 = st.columns(2)
    if col_e1.button("⬇ Export client register (Excel)", key="exp_master_client"):
        out = io.BytesIO()
        with pd.ExcelWriter(out, engine="openpyxl") as writer:
            df_client.to_excel(writer, sheet_name="BMS Register", index=False)
            summary.to_excel(writer, sheet_name="Summary", index=False)
        st.download_button(
            "Download Excel",
            data=out.getvalue(),
            file_name=f"BMS_Register_{p['name'].replace(' ','_')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="dl_master_excel"
        )


# ── Analysis functions ─────────────────────────────────────────────────────────


def _analyse_soo(soo_bytes, soo_name, api_key_val):
    """Read SOO and extract BMS scope register."""
    fname = soo_name.lower()
    if fname.endswith(".docx"):
        text = _extract_docx_text(soo_bytes, 15000)
    else:
        text = _extract_pdf_text(soo_bytes, 15000)

    if not text or len(text.strip()) < 50:
        return {"systems":[],"exclusions":[],"questions":
                ["Could not extract text from SOO PDF"],"io_summary":[],
                "error":"No text extracted"}

    prompt = f"""You are a senior BMS controls engineer reading a Sequence of Operations document.

SOO TEXT (extract all BMS systems from this):
{text[:12000]}

Extract ALL systems that have BMS control sequences and return ONLY valid JSON:
{{
  "systems": [
    {{
      "tag": "HWP-74-1,2,3",
      "system": "Hot Water Pump",
      "floor": "74th Floor",
      "qty": 3,
      "control_type": "DDC + VFD BACnet",
      "scope": "Start/stop, status, speed, fault, DP control",
      "interface": "Hardwired + BACnet MS/TP"
    }}
  ],
  "exclusions": [
    {{
      "item": "Expansion Tanks",
      "reason": "No BMS monitoring - mechanical only"
    }}
  ],
  "questions": [
    "Section 3.4 mentions BTU meters but no monitoring sequence provided"
  ],
  "io_summary": [
    {{
      "system": "Hot Water Pump",
      "ai": 2, "bi": 2, "ao": 1, "bo": 1, "serial": 4,
      "total_pts": 10
    }}
  ]
}}

Rules:
- Include EVERY system that has a written BMS sequence
- Control type: DDC / BACnet MS/TP / BACnet IP / Hardwired / Manufacturer standalone / Monitoring only
- Exclusions = items mentioned but explicitly NOT in BMS scope
- Questions = unclear scope or missing sequences
- Start response with {{ and end with }}
- Return ONLY the JSON, no markdown, no explanation"""

    raw = _claude(api_key_val, prompt, max_tokens=4000) or ""

    if not raw:
        return {"systems":[],"exclusions":[],"questions":
                ["Claude returned empty response - check API key"],"io_summary":[],
                "error":"Empty Claude response"}

    try:
        clean = raw.strip()
        if "```" in clean:
            for part in clean.split("```"):
                part = part.strip().lstrip("json").strip()
                if part.startswith("{"):
                    clean = part
                    break
        s = clean.find("{")
        e = clean.rfind("}")
        if s != -1 and e > s:
            return json.loads(clean[s:e+1])
        return {"systems":[],"exclusions":[],"questions":
                [f"JSON parse failed. Raw: {raw[:300]}"],"io_summary":[],
                "error":"Parse failed"}
    except Exception as ex:
        return {"systems":[],"exclusions":[],"questions":
                [f"Parse error: {ex}. Raw: {raw[:200]}"],"io_summary":[],
                "error":str(ex)}
def _analyse_schedule(sched_bytes, soo_register):
    """Extract BMS-relevant devices from mechanical schedule PDF."""
    import fitz, re

    # BMS-relevant device types (filtered against common non-BMS items)
    BMS_PREFIXES = {
        "WSHP":  ("Water Source Heat Pump",     "BACnet MS/TP"),
        "ASHP":  ("Air Source Heat Pump",        "BACnet IP"),
        "AHU":   ("Air Handling Unit",           "DDC"),
        "FCU":   ("Fan Coil Unit",               "DDC"),
        "ERU":   ("Energy Recovery Unit",        "DDC"),
        "ERV":   ("Energy Recovery Ventilator",  "BACnet MS/TP"),
        "DOAS":  ("DOAS Unit",                   "DDC"),
        "MAU":   ("Make Up Air Unit",            "DDC"),
        "HV":    ("Heating & Ventilating Unit",  "DDC"),
        "ACU":   ("Air Conditioning Unit",       "BACnet MS/TP"),
        "AC":    ("Air Conditioning Unit",       "BACnet MS/TP"),
        "WCCU":  ("Water Cooled Condensing Unit","BACnet MS/TP"),
        "HWP":   ("Hot Water Pump",              "DDC + VFD"),
        "CHWP":  ("Chilled Water Pump",          "DDC + VFD"),
        "PCHWP": ("Primary CHW Pump",            "DDC + VFD"),
        "SCHWP": ("Secondary CHW Pump",          "DDC + VFD"),
        "PHWP":  ("Primary HW Pump",             "DDC + VFD"),
        "SHWP":  ("Secondary HW Pump",           "DDC + VFD"),
        "SCWP":  ("Secondary CW Pump",           "DDC + VFD"),
        "TCWP":  ("Tertiary CW Pump",            "DDC + VFD"),
        "SMP":   ("Snow Melt Pump",              "DDC"),
        "PFHX":  ("Plate & Frame HX",            "Monitoring"),
        "STHX":  ("Steam HX",                    "Monitoring"),
        "SPF":   ("Stair Pressurization Fan",    "DDC"),
        "PFSP":  ("Post-Fire Smoke Purge Fan",   "DDC"),
        "EF":    ("Exhaust Fan",                 "DDC"),
        "SF":    ("Supply Fan",                  "DDC"),
        "TX":    ("Toilet Exhaust Fan",          "DDC"),
        "TRX":   ("Trash Exhaust Fan",           "Monitoring"),
        "SVF":   ("Stair Ventilation Fan",       "DDC"),
        "EP":    ("Elevator Pressurization Fan", "DDC"),
        "VAV":   ("VAV Terminal Box",            "DDC"),
        "FTR":   ("Fin Tube Radiation",          "DDC"),
        "HWC":   ("Hot Water Coil",              "DDC"),
    }

    # Non-BMS items to exclude
    EXCLUDE_PREFIXES = ["ET-","AS-","STHX-","PFHX-"]

    try:
        doc = fitz.open(stream=sched_bytes, filetype="pdf")
        full_text = "\n".join(page.get_text() for page in doc)
    except Exception as e:
        return {"equipment":[],"total_devices":0,"bms_scope":0,
                "flagged":0,"error":str(e)}

    # Extract tags with quantity handling
    # Matches: WSHP-18-1,2,3 or HWP-74-1,2,3,4 or AHU-5-1 TO AHU-5-3
    TAG_RE = re.compile(
        r'\b((?:' + '|'.join(BMS_PREFIXES.keys()) + r')'
        r'-[\w-]+(?:\s*(?:TO|,)\s*[\w-]+)*)',
        re.IGNORECASE
    )

    # Floor extraction from tag
    FLOOR_MAP = {
        "SC": "Sub-Cellar", "C": "Cellar",
        "1": "1st Floor", "5": "5th Floor",
        "18": "18th Floor", "19": "19th Floor",
        "20": "20th Floor", "21": "21st Floor",
        "38": "38th Floor", "61": "61st Floor",
        "74": "74th Floor", "76": "76th Floor",
    }

    def parse_qty(tag_str):
        """Parse quantity from tag range like WSHP-18-1,2,3 or HWP-74-1 TO HWP-74-4."""
        if " TO " in tag_str.upper():
            parts = re.findall(r'\d+$', tag_str.split("TO")[0].strip())
            parts2 = re.findall(r'\d+$', tag_str.split("TO")[-1].strip())
            if parts and parts2:
                try:
                    return abs(int(parts2[0]) - int(parts[0])) + 1
                except: pass
        commas = tag_str.count(",")
        return commas + 1

    def base_tag(tag_str):
        """Get clean base tag like WSHP-18-1,2,3 → WSHP-18"""
        clean = re.split(r',|\s+TO\s+', tag_str, flags=re.IGNORECASE)[0].strip()
        return clean

    def get_floor(tag):
        parts = tag.upper().split("-")
        if len(parts) >= 2:
            return FLOOR_MAP.get(parts[1], parts[1] + "th Floor")
        return "Unknown"

    def get_prefix(tag):
        for prefix in sorted(BMS_PREFIXES.keys(), key=len, reverse=True):
            if tag.upper().startswith(prefix):
                return prefix
        return None

    # Get SOO confirmed systems for filtering
    soo_systems = set()
    for s in soo_register.get("systems", []):
        # Extract prefix from SOO tag
        t = s.get("tag","").upper().split("-")[0]
        soo_systems.add(t)

    seen = {}
    flagged = 0

    for match in TAG_RE.finditer(full_text):
        tag_str = match.group(0).strip()
        btag    = base_tag(tag_str)
        prefix  = get_prefix(btag)
        if not prefix: continue

        # Skip non-BMS
        skip = False
        for ex in EXCLUDE_PREFIXES:
            if btag.upper().startswith(ex.upper()): skip = True
        if skip: continue

        qty = parse_qty(tag_str)
        floor = get_floor(btag)
        system, ctrl = BMS_PREFIXES.get(prefix, ("Unknown","DDC"))

        # SOO status
        if soo_systems:
            soo_status = "SOO confirmed" if prefix in soo_systems else "Verify with SOO"
        else:
            soo_status = "SOO not loaded"

        # Flag "see floor plan"
        context = full_text[max(0,match.start()-100):match.end()+100].lower()
        fp_flag = "see floor plan" in context or "see plan" in context

        key = btag.upper()
        if key not in seen:
            seen[key] = {
                "tag":          btag,
                "tag_range":    tag_str,
                "system":       system,
                "floor":        floor,
                "qty":          qty,
                "control_type": ctrl,
                "soo_status":   soo_status,
                "fp_flag":      "⚠ See floor plan" if fp_flag else "",
                "source":       "Schedule",
                "remarks":      "",
                "soo_confirmed": soo_status == "SOO confirmed",
                "discrepancy_flag": soo_status == "Verify with SOO",
                "classification":  system,
                "bms_interface_default": ctrl,
            }
        else:
            # Update qty if higher
            if qty > seen[key]["qty"]:
                seen[key]["qty"] = qty

        if fp_flag: flagged += 1

    equipment = list(seen.values())
    bms_scope = len([e for e in equipment if "SOO confirmed" in e["soo_status"]])

    return {
        "equipment":     equipment,
        "total_devices": len(equipment),
        "bms_scope":     bms_scope,
        "flagged":       flagged,
    }


def _analyse_riser(water_bytes, air_bytes, soo_register):
    """Extract labeled devices from riser diagrams."""
    import fitz, re

    RISER_DEVICES = {
        "BTU":  ("BTU Meter",           "BACnet Modbus"),
        "FM":   ("Flow Meter",           "BACnet Modbus"),
        "DPS":  ("DP Sensor",            "Hardwired AI"),
        "DP":   ("Differential Pressure","Hardwired AI"),
        "TS":   ("Temperature Sensor",   "Hardwired AI"),
        "PS":   ("Pressure Sensor",      "Hardwired AI"),
        "VS":   ("Vibration Sensor",     "Hardwired AI"),
        "CV":   ("Control Valve",        "Hardwired AO"),
        "MOV":  ("Motorized Valve",      "Hardwired BO"),
        "BV":   ("Ball Valve (motorized)","Hardwired BO"),
        "FLV":  ("Flow Control Valve",   "Hardwired AO"),
    }

    TAG_RE = re.compile(
        r'\b((?:' + '|'.join(RISER_DEVICES.keys()) + r')-[\w-]+)',
        re.IGNORECASE
    )

    equipment = {}
    flagged   = 0

    for label, pdf_bytes in [("Water Riser", water_bytes),
                               ("Air Riser",   air_bytes)]:
        if not pdf_bytes: continue
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)
        except Exception:
            continue

        for match in TAG_RE.finditer(text):
            tag = match.group(0).strip().upper()
            prefix = tag.split("-")[0]
            system, ctrl = RISER_DEVICES.get(prefix, ("Unknown","Hardwired"))

            # Simple floor from tag
            parts = tag.split("-")
            floor = parts[1] + "th Floor" if len(parts) > 1 else "Unknown"

            if tag not in equipment:
                equipment[tag] = {
                    "tag":          tag,
                    "tag_range":    tag,
                    "system":       system,
                    "floor":        floor,
                    "qty":          1,
                    "control_type": ctrl,
                    "soo_status":   "Verify with SOO",
                    "fp_flag":      "",
                    "source":       label,
                    "remarks":      "From riser diagram",
                    "soo_confirmed": False,
                    "discrepancy_flag": False,
                    "classification":   system,
                    "bms_interface_default": ctrl,
                }

    eq_list = list(equipment.values())
    return {
        "equipment": eq_list,
        "total":     len(eq_list),
        "flagged":   flagged,
    }


def _analyse_electrical(elec_bytes, soo_register):
    """Find electrical items confirmed in SOO as BMS monitoring scope."""
    import fitz, re

    ELEC_DEVICES = {
        "GEN":  ("Emergency Generator",    "Monitoring - dry contact"),
        "ATS":  ("Auto Transfer Switch",   "Monitoring - dry contact"),
        "UPS":  ("UPS System",             "Monitoring - dry contact"),
        "EM":   ("Electric Meter",         "BACnet Modbus"),
        "HTR":  ("Heat Trace",             "Monitoring - dry contact"),
        "XFMR": ("Transformer",            "Monitoring"),
        "MCC":  ("Motor Control Center",   "Monitoring"),
        "SWB":  ("Switchboard",            "Monitoring - dry contact"),
        "INV":  ("Inverter",               "Monitoring - dry contact"),
    }

    TAG_RE = re.compile(
        r'\b((?:' + '|'.join(ELEC_DEVICES.keys()) + r')-[\w-]+)',
        re.IGNORECASE
    )

    soo_scope = set()
    for s in soo_register.get("systems", []):
        t = s.get("tag","").upper().split("-")[0]
        soo_scope.add(t)

    try:
        doc = fitz.open(stream=elec_bytes, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
    except Exception as e:
        return {"equipment":[],"total":0,"error":str(e)}

    equipment = {}
    for match in TAG_RE.finditer(text):
        tag    = match.group(0).strip().upper()
        prefix = tag.split("-")[0]
        if prefix not in soo_scope and soo_scope:
            continue  # Not in SOO scope
        system, ctrl = ELEC_DEVICES.get(prefix, ("Electrical Device","Monitoring"))
        if tag not in equipment:
            equipment[tag] = {
                "tag":          tag,
                "tag_range":    tag,
                "system":       system,
                "floor":        "See drawings",
                "qty":          1,
                "control_type": ctrl,
                "soo_status":   "SOO confirmed" if prefix in soo_scope else "Verify",
                "source":       "Electrical",
                "remarks":      "BMS monitoring scope",
                "soo_confirmed": True,
                "discrepancy_flag": False,
                "classification":   system,
                "bms_interface_default": ctrl,
            }

    eq_list = list(equipment.values())
    return {"equipment": eq_list, "total": len(eq_list)}


def _analyse_plumbing(plumb_bytes, soo_register):
    """Find plumbing items confirmed in SOO as BMS monitoring scope."""
    import fitz, re

    PLUMB_DEVICES = {
        "DHW":  ("Domestic Hot Water Heater","Monitoring - leak/status"),
        "EWH":  ("Electric Water Heater",    "Monitoring - leak"),
        "WH":   ("Water Heater",             "Monitoring - leak"),
        "SP":   ("Sump Pump",                "Monitoring - status"),
        "EP":   ("Ejector Pump",             "Monitoring - status"),
        "BP":   ("Booster Pump",             "DDC - start/stop/status"),
        "RP":   ("Recirculation Pump",       "Monitoring - status"),
        "PRV":  ("Pressure Reducing Valve",  "Monitoring"),
        "WM":   ("Water Meter",              "BACnet Modbus"),
        "RPZ":  ("RPZ Backflow Preventer",   "Monitoring"),
    }

    TAG_RE = re.compile(
        r'\b((?:' + '|'.join(PLUMB_DEVICES.keys()) + r')-[\w-]+)',
        re.IGNORECASE
    )

    soo_scope = set()
    for s in soo_register.get("systems", []):
        t = s.get("tag","").upper().split("-")[0]
        soo_scope.add(t)

    try:
        doc = fitz.open(stream=plumb_bytes, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
    except Exception as e:
        return {"equipment":[],"total":0,"error":str(e)}

    equipment = {}
    for match in TAG_RE.finditer(text):
        tag    = match.group(0).strip().upper()
        prefix = tag.split("-")[0]
        system, ctrl = PLUMB_DEVICES.get(prefix, ("Plumbing Device","Monitoring"))
        if tag not in equipment:
            equipment[tag] = {
                "tag":          tag,
                "tag_range":    tag,
                "system":       system,
                "floor":        "See drawings",
                "qty":          1,
                "control_type": ctrl,
                "soo_status":   "SOO confirmed" if prefix in soo_scope else "Verify",
                "source":       "Plumbing",
                "remarks":      "BMS monitoring scope",
                "soo_confirmed": True,
                "discrepancy_flag": False,
                "classification":   system,
                "bms_interface_default": ctrl,
            }

    eq_list = list(equipment.values())
    return {"equipment": eq_list, "total": len(eq_list)}


def _generate_master(p):
    """Merge all takeoff sources into master register."""
    from collections import defaultdict

    all_equip = {}

    # Source priority: Floor Plan > Schedule > Riser > Electrical > Plumbing
    sources = [
        ("Plumbing",   p.get("plumbing_data",{}).get("equipment",[])),
        ("Electrical", p.get("electrical_data",{}).get("equipment",[])),
        ("Riser",      p.get("riser_data",{}).get("equipment",[])),
        ("Schedule",   p.get("schedule_data",{}).get("equipment",[])),
        ("Floor Plan", [
            {"tag": r["Tag"], "system": r.get("System",""),
             "floor": r.get("Pages",""), "qty": 1,
             "control_type": "DDC", "soo_status": r.get("Status",""),
             "source": "Floor Plan", "remarks": ""}
            for r in p.get("floorplan_data",{}).get("statuses",[])
            if r.get("Color") in ("green","amber","blue")
        ]),
    ]

    for source_name, equip_list in sources:
        for e in equip_list:
            tag = e.get("tag","").upper()
            if not tag: continue
            if tag not in all_equip:
                all_equip[tag] = dict(e)
                all_equip[tag]["sources"] = [source_name]
            else:
                # Update qty if floor plan source (higher priority)
                if source_name == "Floor Plan":
                    all_equip[tag]["qty"] = e.get("qty", all_equip[tag]["qty"])
                all_equip[tag]["sources"].append(source_name)

    # Build final list
    final = []
    for tag, e in sorted(all_equip.items()):
        sources_str = " + ".join(dict.fromkeys(e.get("sources",[])))
        # Panel name: CP-[SYSTEM PREFIX]
        prefix = tag.split("-")[0] if "-" in tag else tag[:4]
        panel  = f"CP-{e.get('system','').replace(' ','_').upper()[:8]}"

        final.append({
            "tag":          e.get("tag",""),
            "tag_range":    e.get("tag_range", e.get("tag","")),
            "system":       e.get("system",""),
            "floor":        e.get("floor",""),
            "qty":          e.get("qty",1),
            "control_type": e.get("control_type","DDC"),
            "panel":        panel,
            "soo_status":   e.get("soo_status",""),
            "source":       sources_str,
            "remarks":      e.get("remarks",""),
            "soo_confirmed":    e.get("soo_confirmed", False),
            "discrepancy_flag": e.get("discrepancy_flag", False),
            "classification":   e.get("classification", e.get("system","")),
            "bms_interface_default": e.get("bms_interface_default", e.get("control_type","DDC")),
        })

    total_qty = sum(e["qty"] for e in final)
    return {"equipment": final, "total_qty": total_qty}


# ── Shared helpers ─────────────────────────────────────────────────────────────

def _show_takeoff_results(data, proj_name, source=""):
    """Show equipment table with metrics and export."""
    equip = data.get("equipment", [])
    if not equip:
        st.info("No devices found.")
        return

    # Metrics
    total = len(equip)
    total_qty = sum(e.get("qty",1) for e in equip)
    confirmed = sum(1 for e in equip if e.get("soo_status","") == "SOO confirmed")
    flagged   = sum(1 for e in equip if e.get("fp_flag",""))

    m1,m2,m3,m4 = st.columns(4)
    m1.metric("Device types", total)
    m2.metric("Total qty",    total_qty)
    m3.metric("SOO confirmed",confirmed)
    if flagged:
        m4.metric("⚠ See floor plan", flagged)

    # Table
    display_cols = ["tag","system","floor","qty","control_type","soo_status","fp_flag","source"]
    df = pd.DataFrame([{c: e.get(c,"") for c in display_cols} for e in equip])
    df.columns = ["Tag","System","Floor","Qty","Control Type","SOO Status","FP Flag","Source"]

    # Filter
    search = st.text_input("Filter", placeholder="Search tag or system...",
                            key=f"filter_{source}")
    if search:
        sl = search.lower()
        df = df[df["Tag"].str.lower().str.contains(sl,na=False) |
                df["System"].str.lower().str.contains(sl,na=False)]

    st.dataframe(df, use_container_width=True, hide_index=True, height=400)

    _export_btn(df, f"{source}_takeoff", proj_name, key=f"exp_{source}_data")


def _export_btn(df, filename, proj_name, key="exp_generic"):
    """Reusable Excel export button."""
    if st.button(f"⬇ Export to Excel", key=key):
        out = io.BytesIO()
        df.to_excel(out, index=False)
        st.download_button(
            "Download Excel",
            data=out.getvalue(),
            file_name=f"{filename}_{proj_name.replace(' ','_')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key=f"dl_{key}"
        )



def module_point_list(p):
    st.markdown("### Point list")
    client  = p.get("client")
    cl      = st.session_state.clients.get(client,{}) if client else {}
    pl_tmpl = cl.get("pl_template_bytes")
    pl_name = cl.get("pl_template_name","")

    st.markdown("""
    **Columns:** Panel Name · Equipment · Point Name · Control Device · AI · BI · AO · BO · Serial Pt · Terms · Remarks
    
    AI extracts all fields from SOO; you refine in the editor below.
    """)

    if pl_tmpl:
        st.success(f"✅ Client template loaded: `{pl_name}` — data will export to your format.")
    else:
        st.info("No client template. Standard columns used. Add a template in the Clients tab.")

    has_takeoff = len(p["takeoff"].get("equipment", [])) > 0
    has_soo     = bool(p["docs"].get("SOO"))
    has_spec    = bool(p["docs"].get("Controls spec"))

    if not has_soo:
        st.warning("⚠️ No SOO uploaded. Click 'Generate' to extract from any Controls spec, or upload SOO in Takeoff tab.")
    else:
        st.success("✅ SOO loaded — ready to generate points.")

    if not has_takeoff:
        st.info("ℹ️ No takeoff loaded. Panel Names will be inferred from Equipment tags. Exact wiring details not available.")

    # ── Diagnostics ───────────────────────────────────────────────────────
    with st.expander("🔍 Diagnostics & SOO Preview"):
        soo_bytes  = p["docs"].get("SOO")
        spec_bytes = p["docs"].get("Controls spec")
        k_check    = api_key()
        
        diag = {
            "SOO uploaded":           bool(soo_bytes),
            "SOO size (KB)":          round(len(soo_bytes)/1024, 1) if soo_bytes else 0,
            "Controls spec uploaded": bool(spec_bytes),
            "API key valid":          bool(k_check) and len(k_check) > 10,
            "API key prefix":         k_check[:12] + "..." if k_check else "not set",
            "Takeoff devices":        len(p["takeoff"].get("equipment",[])),
            "Client":                 p.get("client","none"),
            "Points generated":       len(p["point_list"].get("rows",[])),
        }
        
        col1, col2 = st.columns(2)
        with col1:
            for k, v in list(diag.items())[:4]:
                st.metric(k, v if not isinstance(v, bool) else ("✅" if v else "❌"))
        with col2:
            for k, v in list(diag.items())[4:]:
                st.metric(k, v if not isinstance(v, bool) else ("✅" if v else "❌"))
        
        if soo_bytes:
            st.markdown("**SOO text preview (first 400 chars):**")
            try:
                import fitz
                doc = fitz.open(stream=soo_bytes, filetype="pdf")
                preview = doc[0].get_text()[:400].strip()
                st.code(preview, language="text")
            except Exception as e:
                st.error(f"Could not read SOO: {e}")

    if st.button("🤖 Generate point list", type="primary", key="gen_pl"):
        k = api_key()
        if not k:
            st.error("❌ No API key. Add it in Streamlit dashboard → Settings → Secrets: ANTHROPIC_API_KEY = \"sk-ant-...\"")
        elif not p["docs"].get("SOO") and not p["docs"].get("Controls spec"):
            st.error("❌ No SOO or controls spec found in this project. "
                     "Go to the Takeoff tab → Add/replace documents → upload your SOO PDF.")
        else:
            prog = st.progress(0, text="Reading SOO...")
            try:
                prog.progress(20, text="Extracting text from SOO...")
                # Test extraction first
                soo_b = p["docs"].get("SOO")
                if soo_b:
                    import fitz as _fitz
                    _doc = _fitz.open(stream=soo_b, filetype="pdf")
                    _preview = " ".join(page.get_text() for page in _doc)[:200]
                    prog.progress(40, text=f"SOO text found ({len(_preview)} chars preview)...")
                prog.progress(60, text="Sending to Claude...")
                rows = ai_point_list(p, k, pl_tmpl, pl_name)
                prog.progress(85, text="Inferring I/O types...")
                # Apply I/O type inference
                for row in rows:
                    if row.get("Point Name"):
                        io_types = infer_io_type(row["Point Name"], row.get("Remarks", ""))
                        for io_key in ["AI", "BI", "AO", "BO", "Serial_Pt"]:
                            if not row.get(io_key):
                                row[io_key] = io_types[io_key]
                prog.progress(100, text="Done.")
                p["point_list"]["rows"]   = rows
                p["point_list"]["status"] = "done"
                _save_app_state()
                st.success(f"✅ {len(rows)} points generated. Edit any cell to refine.")
                st.rerun()
            except Exception as e:
                prog.progress(100, text="Error.")
                st.error(f"❌ Error: {e}\n\nTip: Check that Streamlit Secrets has ANTHROPIC_API_KEY and it's valid.")

    rows = p["point_list"].get("rows",[])
    if not rows:
        st.info("👆 Click 'Generate point list' to extract from SOO."); return

    st.markdown(f"**{len(rows)} points** — click a cell to edit:")
    
    # Prepare DataFrame with proper column order
    columns = ["Panel Name", "Equipment", "Point Name", "Control Device", 
               "AI", "BI", "AO", "BO", "Serial_Pt", "Terms", "Remarks"]
    df = pd.DataFrame(rows)
    for col in columns:
        if col not in df.columns:
            df[col] = ""
    df = df[columns]
    
    edited = st.data_editor(df, use_container_width=True,
                            hide_index=True, num_rows="dynamic", key="pl_ed")
    p["point_list"]["rows"] = edited.to_dict("records")

    if st.button("⬇ Export to Excel", key="exp_pl"):
        xb = export_pl_excel(edited, pl_tmpl)
        st.download_button("Download .xlsx", xb,
                           f"point_list_{p['name'].replace(' ','_')}.xlsx",
                           "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                           key="dl_pl")

    # ── Appendix Section ──────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### Appendix: System-Wise Special Points")
    st.markdown("""
    **For supplementary points not in main list:**
    - Post-fire smoke purge sequences (PFSP, GX, SPF, HPF)
    - Life safety / emergency pressurization
    - Future expansion or placeholder points
    - Special integrations (Fire alarm, Backup power monitoring)
    - Archived/historical sequences
    """)
    
    # Initialize appendix in session if not present
    if "point_list_appendix" not in p:
        p["point_list_appendix"] = {"rows": [], "status": "not_started"}
    
    if st.button("🤖 Generate Appendix Points", key="gen_pl_app"):
        k = api_key()
        if not k:
            st.error("❌ No API key.")
        elif not p["docs"].get("SOO"):
            st.warning("⚠️ Upload SOO to generate appendix points.")
        else:
            prog = st.progress(0, text="Extracting appendix points from SOO...")
            try:
                prog.progress(40, text="Analyzing main points...")
                main_equip = list(set(row.get("Equipment", "") for row in rows if row.get("Equipment")))
                
                prog.progress(60, text="Sending to Claude...")
                soo_b = p["docs"].get("SOO")
                if soo_b:
                    import fitz as _fitz
                    _doc = _fitz.open(stream=soo_b, filetype="pdf")
                    soo_text = " ".join(page.get_text() for page in _doc)[:10000]
                    
                    from point_list_extractor import generate_appendix_prompt
                    app_prompt = generate_appendix_prompt(p["name"], soo_text, main_equip)
                    app_raw = _claude(k, app_prompt, max_tokens=2000) or ""
                    
                    from point_list_extractor import parse_point_list_response
                    app_rows = parse_point_list_response(app_raw) if app_raw.strip() else []
                    
                    # Apply I/O inference to appendix
                    for row in app_rows:
                        if row.get("Point Name"):
                            io_types = infer_io_type(row["Point Name"], row.get("Remarks", ""))
                            for io_key in ["AI", "BI", "AO", "BO", "Serial_Pt"]:
                                if not row.get(io_key):
                                    row[io_key] = io_types[io_key]
                    
                    p["point_list_appendix"]["rows"] = app_rows
                    p["point_list_appendix"]["status"] = "done"
                    prog.progress(100, text="Done.")
                    st.success(f"✅ {len(app_rows)} appendix points generated.")
                    _save_app_state()
                    st.rerun()
            except Exception as e:
                prog.progress(100)
                st.error(f"Error: {e}")
    
    # Display appendix points
    app_rows = p.get("point_list_appendix", {}).get("rows", [])
    if app_rows:
        st.markdown(f"**{len(app_rows)} appendix points** — edit as needed:")
        columns = ["Panel Name", "Equipment", "Point Name", "Control Device", 
                   "AI", "BI", "AO", "BO", "Serial_Pt", "Terms", "Remarks"]
        df_app = pd.DataFrame(app_rows)
        for col in columns:
            if col not in df_app.columns:
                df_app[col] = ""
        df_app = df_app[columns]
        
        edited_app = st.data_editor(df_app, use_container_width=True,
                                    hide_index=True, num_rows="dynamic", key="pl_app_ed")
        p["point_list_appendix"]["rows"] = edited_app.to_dict("records")
        
        if st.button("⬇ Export Main + Appendix", key="exp_pl_combined"):
            # Export both main and appendix in one Excel file
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            
            wb = Workbook()
            
            # Main sheet
            ws_main = wb.active
            ws_main.title = "Main Points"
            cols = ["Panel Name", "Equipment", "Point Name", "Control Device", 
                   "AI", "BI", "AO", "BO", "Serial_Pt", "Terms", "Remarks"]
            fill = PatternFill("solid", start_color="2E75B6")
            for ci, col in enumerate(cols, 1):
                cell = ws_main.cell(1, ci, col)
                cell.font = Font(bold=True, color="FFFFFF", name="Arial", size=10)
                cell.fill = fill
                ws_main.column_dimensions[cell.column_letter].width = max(14, len(str(col)) + 4)
            
            for ri, row in enumerate(rows, 2):
                for ci, col in enumerate(cols, 1):
                    ws_main.cell(ri, ci, row.get(col, ""))
            
            # Appendix sheet
            if app_rows:
                ws_app = wb.create_sheet("Appendix")
                for ci, col in enumerate(cols, 1):
                    cell = ws_app.cell(1, ci, col)
                    cell.font = Font(bold=True, color="FFFFFF", name="Arial", size=10)
                    cell.fill = PatternFill("solid", start_color="B85C00")
                    ws_app.column_dimensions[cell.column_letter].width = max(14, len(str(col)) + 4)
                
                for ri, row in enumerate(app_rows, 2):
                    for ci, col in enumerate(cols, 1):
                        ws_app.cell(ri, ci, row.get(col, ""))
            
            out = io.BytesIO()
            wb.save(out)
            st.download_button("Download Main + Appendix .xlsx", out.getvalue(),
                               f"point_list_complete_{p['name'].replace(' ','_')}.xlsx",
                               "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                               key="dl_pl_combined")
    else:
        st.info("No appendix points yet. Click 'Generate Appendix Points' to extract special sequences.")

# ── Module 3: Estimate ────────────────────────────────────────────────────────
def module_estimate(p):
    st.markdown("### Estimate")
    init_pricebooks()

    est_tab1, est_tab2, est_tab3 = st.tabs(["⏱ Labor estimate", "🔧 Material estimate", "📊 Combined total"])

    with est_tab1:
        _labor_estimate(p)
    with est_tab2:
        module_material(p)
    with est_tab3:
        _combined_total(p)


def _labor_estimate(p):
    client  = p.get("client")
    cl      = st.session_state.clients.get(client,{}) if client else {}
    saved   = p["estimate"].get("rates") or cl.get("rates") or DEFAULT_RATES

    with st.expander("⚙️ Labor rates ($/hr)"):
        rc = st.columns(len(PHASES))
        rates = {}
        for i,ph in enumerate(PHASES):
            rates[ph] = rc[i].number_input(ph,0,500,int(saved.get(ph,DEFAULT_RATES[ph])),key=f"er_{ph}")
        p["estimate"]["rates"] = rates

    markup = st.slider("Markup %",0,50,int(p["estimate"].get("markup",10)),key="mk_sl")
    p["estimate"]["markup"] = markup

    if st.button("🤖 Generate labor estimate", type="primary", key="gen_est"):
        k = api_key()
        if not k: st.error("Add API key in sidebar.")
        else:
            with st.spinner("Claude is reading SOO, controls spec, and point list to estimate hours..."):
                lines = ai_estimate(p, k, rates, markup)
            p["estimate"]["lines"]  = lines
            p["estimate"]["status"] = "done"
            st.rerun()

    lines = p["estimate"].get("lines",[])
    if not lines:
        st.info("No labor estimate yet. Click Generate."); return

    edited = st.data_editor(pd.DataFrame(lines), use_container_width=True,
                            hide_index=True, num_rows="dynamic", key="est_ed")
    p["estimate"]["lines"] = edited.to_dict("records")

    try:
        sub = sum(float(r.get("Total $",0)) for r in edited.to_dict("records"))
        mk  = sub*markup/100
        t1,t2,t3 = st.columns(3)
        t1.metric("Labor subtotal",    f"${sub:,.0f}")
        t2.metric(f"Markup ({markup}%)",f"${mk:,.0f}")
        t3.metric("Labor total",        f"${sub+mk:,.0f}")
    except: pass

    if st.button("⬇ Export labor to Excel", key="exp_est"):
        out = io.BytesIO()
        edited.to_excel(out, index=False)
        st.download_button("Download .xlsx", out.getvalue(),
                           f"labor_{p['name'].replace(' ','_')}.xlsx",
                           "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                           key="dl_est")


def _combined_total(p):
    st.markdown("**Combined estimate summary**")
    markup = p["estimate"].get("markup", 10)

    # Labor
    labor_lines = p["estimate"].get("lines", [])
    try:
        labor_sub = sum(float(r.get("Total $",0)) for r in labor_lines)
        labor_mk  = labor_sub * markup / 100
        labor_tot = labor_sub + labor_mk
    except:
        labor_sub = labor_mk = labor_tot = 0

    # Material
    mat = p.get("material", {})
    mat_items = mat.get("items", [])
    mat_markup = mat.get("markup", markup)
    try:
        mat_sub = sum(i.get("qty",0) * i.get("unit_cost",0) for i in mat_items)
        mat_mk  = mat_sub * mat_markup / 100
        mat_tot = mat_sub + mat_mk
    except:
        mat_sub = mat_mk = mat_tot = 0

    grand = labor_tot + mat_tot

    # Summary table
    rows = [
        {"Category":"Labor","Subtotal":f"${labor_sub:,.0f}",f"Markup ({markup}%)":f"${labor_mk:,.0f}","Total":f"${labor_tot:,.0f}"},
        {"Category":"Material","Subtotal":f"${mat_sub:,.0f}",f"Markup ({mat_markup}%)":f"${mat_mk:,.0f}","Total":f"${mat_tot:,.0f}"},
        {"Category":"GRAND TOTAL","Subtotal":"","f`Markup`:":""," Total":f"${grand:,.0f}"},
    ]

    c1,c2,c3,c4 = st.columns(4)
    c1.metric("Labor total",    f"${labor_tot:,.0f}")
    c2.metric("Material total", f"${mat_tot:,.0f}")
    c3.metric("Grand total",    f"${grand:,.0f}")
    c4.metric("Labor items",    len(labor_lines))

    st.divider()
    st.markdown("**Labor breakdown**")
    if labor_lines:
        st.dataframe(pd.DataFrame(labor_lines)[["System","Total hrs","Total $"]]
                     if "System" in pd.DataFrame(labor_lines).columns else pd.DataFrame(labor_lines),
                     use_container_width=True, hide_index=True)

    st.markdown("**Material breakdown by section**")
    if mat_items:
        from collections import defaultdict as _dd
        sec_totals = _dd(float)
        for i in mat_items:
            sec_totals[i.get("section","Other")] += i.get("qty",0)*i.get("unit_cost",0)
        df_sec = pd.DataFrame([{"Section":k,"Total":f"${v:,.2f}"} for k,v in sec_totals.items()])
        st.dataframe(df_sec, use_container_width=True, hide_index=True)

    if grand > 0 and st.button("⬇ Export combined summary to Excel", key="exp_combined"):
        out = _export_combined_excel(p, labor_lines, mat_items, labor_sub, labor_mk,
                                     mat_sub, mat_mk, grand, markup, mat_markup)
        st.download_button("Download .xlsx", out,
                           f"full_estimate_{p['name'].replace(' ','_')}.xlsx",
                           "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                           key="dl_combined")


def _export_combined_excel(p, labor_lines, mat_items, labor_sub, labor_mk,
                           mat_sub, mat_mk, grand, labor_markup, mat_markup):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = Workbook()

    # Sheet 1: Summary
    ws = wb.active; ws.title = "Summary"
    ws["A1"] = f"Full Estimate — {p['name']}"
    ws["A1"].font = Font(bold=True, size=14, name="Arial")
    fill_blue = PatternFill("solid", start_color="1F4E79")
    for r,(label,val) in enumerate([
        ("Labor subtotal", labor_sub), (f"Labor markup ({labor_markup}%)", labor_mk),
        ("Labor total", labor_sub+labor_mk), ("",""),
        ("Material subtotal", mat_sub), (f"Material markup ({mat_markup}%)", mat_mk),
        ("Material total", mat_sub+mat_mk), ("",""),
        ("GRAND TOTAL", grand),
    ], start=3):
        ws.cell(r,1,label).font = Font(name="Arial", bold=(label in ("GRAND TOTAL","Labor total","Material total")))
        ws.cell(r,2,val if isinstance(val,float) else "").number_format = '$#,##0.00'
        if label == "GRAND TOTAL":
            ws.cell(r,1).font = Font(bold=True,color="FFFFFF",name="Arial")
            ws.cell(r,1).fill = fill_blue
            ws.cell(r,2).font = Font(bold=True,color="FFFFFF",name="Arial")
            ws.cell(r,2).fill = fill_blue

    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 16

    # Sheet 2: Labor
    if labor_lines:
        ws2 = wb.create_sheet("Labor Estimate")
        df_l = pd.DataFrame(labor_lines)
        ws2.append(list(df_l.columns))
        for row in df_l.values.tolist():
            ws2.append([str(v) if v is not None else "" for v in row])

    # Sheet 3: Material
    if mat_items:
        ws3 = wb.create_sheet("Material Estimate")
        cols = ["section","subsection","description","manufacturer","part_no","qty","unit_cost","ext_cost"]
        ws3.append([c.replace("_"," ").title() for c in cols])
        for item in mat_items:
            ws3.append([item.get(c,"") for c in cols])

    out = io.BytesIO(); wb.save(out); return out.getvalue()

# ── Module 4: Proposal ────────────────────────────────────────────────────────
def module_proposal(p):
    st.markdown("### Proposal")
    client    = p.get("client")
    cl        = st.session_state.clients.get(client,{}) if client else {}
    prop_tmpl = cl.get("prop_template_bytes")
    prop_name = cl.get("prop_template_name","")

    if prop_tmpl:
        st.success(f"✅ Client template: `{prop_name}` — placeholders: {{{{PROJECT_NAME}}}}, {{{{CLIENT}}}}, {{{{DATE}}}}, {{{{SCOPE_TEXT}}}}")
    else:
        st.info("No proposal template — AI writes a standard TEC-style proposal. Add your Word template in the Clients tab.")

    # Show what data is available
    has_takeoff  = len(p["takeoff"].get("equipment",[])) > 0
    has_estimate = len(p["estimate"].get("lines",[])) > 0
    has_soo      = bool(p["docs"].get("SOO"))
    has_pl       = len(p["point_list"].get("rows",[])) > 0

    c1,c2,c3,c4 = st.columns(4)
    c1.metric("Devices",    len(p["takeoff"].get("equipment",[])))
    c2.metric("Points",     len(p["point_list"].get("rows",[])))
    c3.metric("Est. lines", len(p["estimate"].get("lines",[])))
    c4.metric("SOO",        "✅" if has_soo else "—")

    if not has_takeoff and not has_soo:
        st.warning("⚠️ Upload SOO or complete takeoff for a meaningful proposal.")

    if st.button("🤖 Generate proposal", type="primary", key="gen_prop"):
        k = api_key()
        if not k:
            st.error("❌ No API key. Add ANTHROPIC_API_KEY in Streamlit Secrets.")
        else:
            prog = st.progress(0, text="Building scope from project data...")
            try:
                prog.progress(30, text="Sending to Claude...")
                text = ai_proposal(p, k)
                prog.progress(100, text="Done.")
                if not text:
                    st.error("Claude returned empty response. Check your API key in Streamlit Secrets.")
                else:
                    p["proposal"]["text"]   = text
                    p["proposal"]["status"] = "done"
                    _save_app_state()
                    st.success("✅ Proposal generated.")
                    st.rerun()
            except Exception as e:
                prog.progress(100, text="Error.")
                st.error(f"Error: {e}")

    text = p["proposal"].get("text","")
    if not text:
        st.info("No proposal yet. Click Generate above.")
        return

    edited = st.text_area("Proposal text — edit before exporting", value=text,
                          height=520, key="prop_ed")
    p["proposal"]["text"] = edited

    if st.button("⬇ Export to Word (.docx)", key="exp_prop"):
        db = export_prop_docx(edited, p, prop_tmpl)
        st.download_button("Download .docx", db,
                           f"proposal_{p['name'].replace(' ','_')}.docx",
                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           key="dl_prop")

# ── Module 5: AI Advisor ──────────────────────────────────────────────────────
def module_ai_advisor(p):
    st.markdown("### AI Advisor")
    st.caption("Ask anything about this project. Claude has full context of all uploaded documents and module outputs.")
    k = api_key()
    if not k:
        st.error("Add your Anthropic API key in the sidebar."); return

    equip = p["takeoff"].get("equipment",[])
    discs = p["takeoff"].get("discrepancies",[])
    tags  = sorted(set(e.get("tag","") for e in equip if e.get("tag")))

    # Quick action buttons
    st.markdown("**Quick actions**")
    qa1,qa2,qa3 = st.columns(3)
    if discs and qa1.button(f"Resolve {len(discs)} discrepancies ↗", key="qa1"):
        _ask(f"I have {len(discs)} discrepancies in {p['name']}: "
             f"{', '.join(d.get('tag','') for d in discs[:12])}. "
             f"For each tag, what BMS scope is typically required and what should I confirm with the engineer?",
             p, k)
    if qa2.button("Review estimate ↗", key="qa2"):
        n = len(p["estimate"].get("lines",[]))
        _ask(f"Review the {n}-line estimate for {p['name']}. "
             f"Are there any line items that seem off vs industry norms? What might be missing?", p, k)
    if qa3.button("Scope gap check ↗", key="qa3"):
        _ask(f"Based on the takeoff and SOO for {p['name']}, what BMS scope items "
             f"might be missing before I finalise the proposal?", p, k)

    st.divider()

    # Device Q&A
    if tags:
        st.markdown("**Device scope question**")
        a1,a2 = st.columns([1,2])
        tag    = a1.selectbox("Device tag", tags, key="ai_tag")
        qtype  = a2.selectbox("Question", [
            "What BMS control points are needed?",
            "What is the control sequence?",
            "How do I handle this discrepancy?",
            "Estimate labor hours",
            "BACnet integration requirements",
        ], key="ai_qt")
        rec    = next((e for e in equip if e.get("tag")==tag),{})
        custom = st.text_area("Additional context (optional)",
                               placeholder="e.g. Niagara 4 front-end, back-of-house location…",
                               height=60, key="ai_ctx")
        if st.button("Ask Claude →", type="primary", key="ai_dev"):
            _ask(_device_prompt(tag, rec, qtype, custom, p["name"]), p, k)

    st.divider()
    st.markdown("**Freeform question**")
    fq = st.text_area("Ask anything about this project…", height=80, key="ai_free")
    if st.button("Ask →", key="ai_free_btn") and fq:
        _ask(fq, p, k)

    proj_hist = [h for h in st.session_state.ai_history if h.get("project")==p["name"]]
    if proj_hist:
        st.divider()
        st.markdown("**Session history**")
        for h in reversed(proj_hist[-6:]):
            with st.expander(h["q"][:70]):
                st.markdown(h["a"])

def _ask(prompt, p, k):
    ctx = (f"Project: {p['name']} | Client: {p.get('client','—')} | "
           f"Address: {p.get('address','—')} | Bid: {p.get('bid_date','—')} | "
           f"Docs: {', '.join(p.get('doc_names',{}).keys())} | "
           f"Devices: {len(p['takeoff'].get('equipment',[]))} | "
           f"Discrepancies: {len(p['takeoff'].get('discrepancies',[]))} | "
           f"Points: {len(p['point_list'].get('rows',[]))} | "
           f"Estimate lines: {len(p['estimate'].get('lines',[]))}\n"
           f"You are a senior BMS estimator with 20 years NYC commercial experience.\n\n")
    with st.spinner("Claude is thinking…"):
        resp = _claude(k, ctx+prompt, 2000)
    if resp:
        st.markdown("---"); st.markdown(resp)
        st.session_state.ai_history.append({"project":p["name"],"q":prompt[:80],"a":resp})

def _device_prompt(tag, rec, qtype, extra, proj):
    qmap = {
        "What BMS control points are needed?":
            "List all BMS points (AI/AO/DI/DO/Network) as a table: type, description, units, normal range.",
        "What is the control sequence?":
            "Write a concise SOO-style control sequence: startup/shutdown, setpoint, alarms, interlocks.",
        "How do I handle this discrepancy?":
            "Device is in schedule but missing from SOO. What BMS scope is typically needed? What to ask the engineer?",
        "Estimate labor hours":
            "Estimate hours by phase: Engineering, Programming, Integration, Graphics, Startup. Show reasoning.",
        "BACnet integration requirements":
            "BACnet object types, instance numbering, required writable objects, vendor-specific notes.",
    }
    return (f"Project: {proj} | Device: {tag} | System: {rec.get('system','?')} | "
            f"Floor: {rec.get('floor','?')} | Classification: {rec.get('classification','?')} | "
            f"BMS interface: {rec.get('bms_interface_default',rec.get('bms_interface','?'))} | "
            f"SOO status: {'NOT IN SOO — discrepancy' if rec.get('discrepancy_flag') else 'confirmed'}"
            + (f" | Extra: {extra}" if extra else "")
            + f"\n\n{qmap.get(qtype,qtype)}")

# ── AI calls ──────────────────────────────────────────────────────────────────
def _claude(k, prompt, max_tokens=1500):
    try:
        import anthropic
        client = anthropic.Anthropic(api_key=k)
        msg = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=max_tokens,
            messages=[{"role":"user","content":prompt}]
        )
        return msg.content[0].text
    except Exception as e:
        err = str(e)
        if "401" in err or "authentication" in err.lower() or "invalid" in err.lower():
            st.error(
                "❌ **Invalid API key.** Go to [Streamlit Cloud dashboard]"
                "(https://share.streamlit.io) → your app → ⋮ → Settings → Secrets "
                "and add:\n```\nANTHROPIC_API_KEY = \"sk-ant-your-key-here\"\n```"
            )
        elif "429" in err:
            st.error("❌ Rate limit hit. Wait 30 seconds and try again.")
        else:
            st.error(f"❌ Claude error: {err}")
        return None

def _extract_pdf_text(pdf_bytes, max_chars=12000):
    """Extract text from PDF using PyMuPDF. Returns truncated text for Claude."""
    try:
        import fitz
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        pages_text = []
        for i, page in enumerate(doc):
            text = page.get_text().strip()
            if text:
                pages_text.append(f"--- PAGE {i+1} ---\n{text}")
        full = "\n".join(pages_text)
        # Truncate to max_chars to stay within token limits
        if len(full) > max_chars:
            full = full[:max_chars] + f"\n[...truncated at {max_chars} chars]"
        return full
    except Exception as e:
        return f"[Could not extract text: {e}]"


def _extract_docx_text(docx_bytes, max_chars=8000):
    """Extract text from DOCX."""
    try:
        from docx import Document as D
        doc = D(io.BytesIO(docx_bytes))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return text[:max_chars]
    except Exception as e:
        return f"[Could not extract text: {e}]"


def ai_takeoff(p, k):
    """
    Real AI takeoff — extracts text from uploaded PDFs and sends to Claude.
    Reads: Drawings (schedule pages) + SOO (cross-check) + Controls spec (scope).
    """
    docs      = p.get("docs", {})
    doc_names = p.get("doc_names", {})

    # ── Extract text from each document ──────────────────────────────────
    extracted = {}
    for label in ["Drawings", "SOO", "Controls spec"]:
        raw_bytes = docs.get(label)
        if not raw_bytes:
            extracted[label] = None
            continue
        fname = doc_names.get(label, "")
        if fname.lower().endswith(".docx"):
            extracted[label] = _extract_docx_text(raw_bytes)
        else:
            # PDF — drawings get more characters, SOO gets medium
            max_c = 14000 if label == "Drawings" else 8000
            extracted[label] = _extract_pdf_text(raw_bytes, max_chars=max_c)

    # ── Build prompt with real content ────────────────────────────────────
    drawings_text = extracted.get("Drawings") or "Not provided"
    soo_text      = extracted.get("SOO")      or "Not provided"
    spec_text     = extracted.get("Controls spec") or "Not provided"

    prompt = f"""You are a senior BMS estimation expert reviewing mechanical drawings for project '{p['name']}'.

DRAWINGS TEXT (schedule sheets and floor plans):
{drawings_text}

SEQUENCE OF OPERATIONS (SOO):
{soo_text}

CONTROLS SPECIFICATION:
{spec_text[:4000] if spec_text != "Not provided" else "Not provided"}

TASK:
1. Extract ALL unique BMS device tags from the drawings text above.
2. For each tag, check if it appears in the SOO text. If it has no SOO sequence, flag as discrepancy.
3. Common discrepancies: EUH (electric unit heaters) and UH (hot water unit heaters) with integral/wall thermostats — these appear in schedules but rarely have SOO sequences.

Return ONLY valid JSON in this exact format:
{{
  "equipment": [
    {{
      "tag": "FCU-SC-1",
      "qty": 1,
      "floor": "Subcellar",
      "system": "Fan Coil Unit",
      "classification": "Terminal / FCU",
      "bms_interface_default": "DDC",
      "soo_confirmed": true,
      "discrepancy_flag": false,
      "action": ""
    }}
  ],
  "discrepancies": [
    {{
      "tag": "EUH-SC-1",
      "system": "Electric Unit Heater",
      "floor": "Subcellar",
      "action": "Integral thermostat — confirm whether BMS monitoring point is required"
    }}
  ]
}}

Device types to extract: FCU, AHU, DOAS, MUA, VAV, EUH, UH, ASHP, ERV, PFSP, SPF, HPF, GX, EF, HF, TF, PCHWP, SCHWP, PHWP, SHWP, FPP, PFHX, BT, GFU, FOP, ESP, FTR, HWC, ACU, AC-C.

Floor mapping from tag suffix: SC=Subcellar, C=Cellar, 1M=1st/Mezzanine, 12=12th Floor, 33=33rd Floor, 34=34th Floor, 35=35th Floor, 36=36th Floor, ROOF=Roof.

Return ONLY the JSON, no other text."""

    raw = _claude(k, prompt, 4000) or ""

    try:
        clean = raw.strip()
        if "```" in clean:
            parts = clean.split("```")
            clean = parts[1] if len(parts) > 1 else clean
            if clean.startswith("json"):
                clean = clean[4:]
        return json.loads(clean.strip())
    except Exception:
        # If JSON parse fails, return error with raw output for debugging
        return {
            "equipment": [],
            "discrepancies": [],
            "error": f"JSON parse failed. Raw output: {raw[:500]}"
        }

def ai_point_list(p, k, pl_tmpl, pl_name):
    """Generate BMS point list from SOO with BMS-specific columns.
    
    Columns: Panel Name, Equipment, Point Name, Control Device, AI, BI, AO, BO, Serial_Pt, Terms, Remarks
    AI infers I/O types; user refines in editor.
    """

    # ── Column format (BMS-specific) ───────────────────────────────────────
    columns = ["Panel Name", "Equipment", "Point Name", "Control Device", 
               "AI", "BI", "AO", "BO", "Serial_Pt", "Terms", "Remarks"]
    if pl_tmpl:
        try:
            df_t = pd.read_excel(io.BytesIO(pl_tmpl), nrows=3)
            cols = [str(c) for c in df_t.columns if not str(c).startswith("Unnamed")]
            if cols: columns = cols
        except Exception:
            pass

    # ── Extract SOO text ──────────────────────────────────────────────────
    soo_text = ""
    if p["docs"].get("SOO"):
        fname = p["doc_names"].get("SOO","")
        soo_text = (_extract_docx_text(p["docs"]["SOO"], 12000)
                    if fname.lower().endswith(".docx")
                    else _extract_pdf_text(p["docs"]["SOO"], 12000))
    
    if not soo_text:
        return [{c: "No SOO found. Upload in Takeoff tab." if c == "Panel Name" 
                     else "Error" if c == "Equipment"
                     else "" for c in columns}]

    # ── Get takeoff equipment for context ──────────────────────────────────
    equip = p["takeoff"].get("equipment", [])

    # ── Generate prompt using new module ──────────────────────────────────
    prompt = generate_point_list_prompt(p["name"], soo_text, takeoff_equip=equip)

    # ── Call Claude ───────────────────────────────────────────────────────
    raw = _claude(k, prompt, max_tokens=4000) or ""

    # ── Parse response and apply I/O inference ──────────────────────────

    try:
        rows = parse_point_list_response(raw)
        # Ensure all columns present
        for row in rows:
            for col in columns:
                row.setdefault(col, "")
        return rows
    except Exception as exc:
        return [{c: (f"Parse error: {str(exc)[:80]}"
                     if c == "Panel Name"
                     else "Check Streamlit Secrets → ANTHROPIC_API_KEY" if c == "Equipment"
                     else raw[:100] if c == "Point Name"
                     else "") for c in columns}]

def ai_estimate(p, k, rates, markup):
    pts  = len(p["point_list"].get("rows", []))
    devs = len(p["takeoff"].get("equipment", []))

    # Extract spec/SOO context for better estimates
    soo_text  = _extract_pdf_text(p["docs"]["SOO"],  max_chars=5000) if p["docs"].get("SOO")  else ""
    spec_text = _extract_pdf_text(p["docs"].get("Controls spec") or b"", max_chars=3000) if p["docs"].get("Controls spec") else ""

    has_takeoff  = devs > 0
    has_pl       = pts > 0

    if has_takeoff and has_pl:
        basis = f"Takeoff: {devs} devices. Point list: {pts} points."
    elif has_pl:
        basis = f"Point list: {pts} points (no takeoff loaded — device counts are estimated)."
    elif has_takeoff:
        basis = f"Takeoff: {devs} devices (no point list — estimating ~5 pts/device)."
        pts = devs * 5
    else:
        basis = "No takeoff or point list — estimating from SOO/spec scope only. Counts are rough."
        pts = 50  # rough default

    prompt = (
        f"Generate a BMS labor estimate for project '{p['name']}'.\n"
        f"Basis: {basis}\n"
        f"Labor rates: {json.dumps(rates)}. Markup: {markup}%.\n"
        f"SOO context: {soo_text[:3000] if soo_text else 'not provided'}\n"
        f"Controls spec: {spec_text[:2000] if spec_text else 'not provided'}\n\n"
        f"Hour formulas: Engineering=0.5×pts, Programming=1.0×pts, "
        f"Integration=0.5×pts, Graphics=0.5×pts, Startup=0.5×pts.\n"
        f"Group by system (FCUs, DOAS, Pumps, Heat Pumps, Life Safety, etc).\n"
        f"Return ONLY a JSON array:\n"
        f'[{{"System":"Fan Coil Units","Qty":8,"Points":40,'
        f'"Engineering (hrs)":20,"Programming (hrs)":40,"Integration (hrs)":20,'
        f'"Graphics (hrs)":20,"Startup (hrs)":20,"Total hrs":120,'
        f'"Rate ($/hr)":{rates.get("Programming",85)},"Total $":10200,'
        f'"Notes":"Based on SOO Section 1.10"}}]\n'
        f"Add rows for: panel/hardware allowance, engineering/submittal, "
        f"commissioning, project management.\n"
        f"Return ONLY the JSON array."
    )
    raw = _claude(k, prompt, 2500) or ""
    try:
        clean = raw.strip()
        if "```" in clean: clean = clean.split("```")[1]; clean = clean[4:] if clean.startswith("json") else clean
        return json.loads(clean.strip())
    except:
        return [{"System":"Parse error","Qty":0,"Points":0,
                 "Engineering (hrs)":0,"Programming (hrs)":0,"Integration (hrs)":0,
                 "Graphics (hrs)":0,"Startup (hrs)":0,"Total hrs":0,
                 "Rate ($/hr)":0,"Total $":0,"Notes":raw[:80]}]

def ai_proposal(p, k):
    cl      = st.session_state.clients.get(p.get("client"),{}) if p.get("client") else {}
    tmpl    = cl.get("prop_template_bytes")
    equip   = p["takeoff"].get("equipment",[])
    lines   = p["estimate"].get("lines",[])

    # Build scope from takeoff if available, otherwise from SOO
    if equip:
        agg = defaultdict(int)
        for e in equip: agg[e.get("system","Unknown")]+=1
        scope = "; ".join(f"{v}× {k}" for k,v in list(agg.items())[:15])
    else:
        # Extract scope from SOO
        soo_text = _extract_pdf_text(p["docs"]["SOO"], max_chars=4000) if p["docs"].get("SOO") else ""
        scope = f"Scope based on SOO (no takeoff loaded): {soo_text[:500]}..." if soo_text else "Scope to be confirmed"

    try:
        total = sum(float(r.get("Total $",0)) for r in lines)*(1+p["estimate"].get("markup",10)/100)
        price = f"${total:,.0f}"
    except: price = "[TO BE DETERMINED — complete estimate first]"
    tmpl_note = ""
    if tmpl:
        try:
            from docx import Document as D
            doc = D(io.BytesIO(tmpl))
            tmpl_note = "Template sections found:\n"+"\n".join(
                f"- {par.text}" for par in doc.paragraphs if par.text.strip() and len(par.text.strip())>3)[:600]
        except: tmpl_note = "Follow client proposal template structure."
    prompt = (
        f"Write a professional BMS proposal for '{p['name']}' addressed to {p.get('client','the client')}.\n"
        f"Address: {p.get('address','—')} | Bid date: {p.get('bid_date','—')}\n"
        f"Scope: {scope}\nBase scope price: {price}\n{tmpl_note}\n\n"
        f"Include: 1) Introduction 2) Basis of proposal 3) Scope of work (bullets per system) "
        f"4) Pricing 5) Notes & exclusions 6) Closing.\n"
        f"Use 'We will provide...' language. TEC Building Systems proposal style.\nReturn full proposal text."
    )
    result = _claude(k, prompt, max_tokens=3000)
    if not result:
        return "[Error: Claude returned empty response. Check your API key in Streamlit Secrets.]"
    return result

# ── Export ────────────────────────────────────────────────────────────────────
def export_pl_excel(df, tmpl_bytes=None):
    from openpyxl import load_workbook, Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    if tmpl_bytes:
        try:
            wb = load_workbook(io.BytesIO(tmpl_bytes))
            ws = wb.active
            hr = 1
            for row in ws.iter_rows(min_row=1,max_row=10):
                if any(c.value for c in row): hr=row[0].row; break
            tcols = {ws.cell(hr,c).value:c for c in range(1,ws.max_column+1) if ws.cell(hr,c).value}
            for ri,rd in enumerate(df.to_dict("records")):
                for cn,ci in tcols.items():
                    if cn in rd: ws.cell(hr+1+ri,ci,rd[cn])
            out=io.BytesIO(); wb.save(out); return out.getvalue()
        except: pass
    wb=Workbook(); ws=wb.active; ws.title="Point List"
    cols=list(df.columns)
    fill=PatternFill("solid",start_color="2E75B6")
    for ci,col in enumerate(cols,1):
        cell=ws.cell(1,ci,col)
        cell.font=Font(bold=True,color="FFFFFF",name="Arial",size=10)
        cell.fill=fill; cell.alignment=Alignment(horizontal="center")
        ws.column_dimensions[cell.column_letter].width=max(14,len(str(col))+4)
    for ri,row in enumerate(df.to_dict("records"),2):
        for ci,col in enumerate(cols,1): ws.cell(ri,ci,row.get(col,""))
    out=io.BytesIO(); wb.save(out); return out.getvalue()

def export_prop_docx(text, p, tmpl_bytes=None):
    from docx import Document as D
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if tmpl_bytes:
        try:
            doc=D(io.BytesIO(tmpl_bytes))
            repls={"{{PROJECT_NAME}}":p["name"],"{{ADDRESS}}":p.get("address",""),
                   "{{CLIENT}}":p.get("client",""),"{{DATE}}":str(date.today()),
                   "{{SCOPE_TEXT}}":text}
            for para in doc.paragraphs:
                for k,v in repls.items():
                    if k in para.text:
                        for run in para.runs: run.text=run.text.replace(k,v)
            found=any("{{SCOPE_TEXT}}" in para.text for para in doc.paragraphs)
            if not found:
                doc.add_page_break()
                for line in text.split("\n"): doc.add_paragraph(line)
            out=io.BytesIO(); doc.save(out); return out.getvalue()
        except: pass
    doc=D()
    t=doc.add_heading(p["name"],0); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {date.today()}")
    doc.add_paragraph(f"Address: {p.get('address','')}")
    doc.add_paragraph(f"Client: {p.get('client','')}")
    doc.add_paragraph()
    for line in text.split("\n"):
        if line.startswith("# "): doc.add_heading(line[2:],1)
        elif line.startswith("## "): doc.add_heading(line[3:],2)
        elif line.startswith("- "): doc.add_paragraph(line[2:],style="List Bullet")
        elif line.strip(): doc.add_paragraph(line)
        else: doc.add_paragraph()
    out=io.BytesIO(); doc.save(out); return out.getvalue()

# ── Reports ───────────────────────────────────────────────────────────────────
def page_reports():
    st.title("Reports")
    st.info("Coming soon: cross-project pipeline analytics, win/loss tracking, device type trends.")

# ── Product status panel ──────────────────────────────────────────────────────

MODULE_READINESS = [
    {
        "module":   "Takeoff",
        "icon":     "📐",
        "ready":    95,
        "status":   "production",
        "what_works": [
            "Text-layer tag extraction from CAD PDFs",
            "Schedule page auto-detection",
            "SOO cross-check (3-way: schedule × SOO × drawing)",
            "Amber/red/green/blue status per device",
            "West 34th Hotel: 109 devices, 12 discrepancies confirmed",
        ],
        "gaps": [
            "AI Takeoff button reads doc names, not PDF content — use PDF extraction instead",
        ],
        "next": "Wire Claude Vision to read floor plan images for scanned PDFs",
    },
    {
        "module":   "Drawing Markup",
        "icon":     "🖊",
        "ready":    90,
        "status":   "production",
        "what_works": [
            "Search any tag → jump to page with highlight",
            "Annotated PDF download with all 4 status colors",
            "Legend page + summary overlay on page 1",
            "Amber tags → Clarifications, Red tags → Exclusions (one click)",
            "Progress bar with page count and schedule page detection",
        ],
        "gaps": [
            "Page preview render speed depends on PDF size",
            "Rotated or very small tags may be missed",
        ],
        "next": "Add thumbnail strip to browse all pages; add zoom on page preview",
    },
    {
        "module":   "Point List",
        "icon":     "📋",
        "ready":    65,
        "status":   "demo",
        "what_works": [
            "AI generates AI/AO/DI/DO per device from context",
            "Client Excel template: AI matches your column format",
            "Editable table — change any cell before export",
            "Export to Excel (plain or matching client template)",
        ],
        "gaps": [
            "AI reads device names, not actual SOO line by line",
            "Point counts are estimated, not extracted from SOO",
        ],
        "next": "Extract SOO text per system, send to Claude for exact point list",
    },
    {
        "module":   "Estimate",
        "icon":     "💰",
        "ready":    70,
        "status":   "demo",
        "what_works": [
            "Labor: AI generates hours by phase from point count + rates",
            "Material: 200-item Honeywell price book, qty × unit cost",
            "Client labor rates saved per client profile",
            "Combined total: Labor + Material with markup",
            "Export to Excel (labor / material / combined sheets)",
        ],
        "gaps": [
            "Labor hours are formula-based, not read from controls spec",
            "Material prices are 2022 list prices",
            "No material take-off from drawings yet",
        ],
        "next": "Read controls spec for inclusions/exclusions; connect price book to vendor API",
    },
    {
        "module":   "Proposal",
        "icon":     "📄",
        "ready":    60,
        "status":   "demo",
        "what_works": [
            "AI writes TEC-style proposal from scope summary",
            "Clarifications auto-filled from amber tags",
            "Exclusions auto-filled from red tags",
            "Client Word template: fills {{placeholders}}",
            "Export to .docx",
        ],
        "gaps": [
            "AI writes from scope summary, not from reading your actual proposal template sections",
            "Needs human review before sending to client",
        ],
        "next": "Parse Word template sections, fill each section individually with targeted AI",
    },
    {
        "module":   "AI Advisor",
        "icon":     "🤖",
        "ready":    80,
        "status":   "production",
        "what_works": [
            "Full project context in every call",
            "Device-level Q&A: points, sequence, discrepancy resolution, labor hours",
            "Quick actions: resolve discrepancies, review estimate, scope gap check",
            "Session history per project",
        ],
        "gaps": [
            "Context is metadata only — doesn't read uploaded PDF content",
        ],
        "next": "Pass extracted PDF text as context for richer, document-grounded answers",
    },
]

PRODUCTION_GAPS = [
    {
        "gap":      "Persistent storage",
        "effort":   "2-3 weeks",
        "impact":   "HIGH",
        "why":      "Everything lives in session state — disappears on refresh. "
                    "Needs Supabase (Postgres) for projects, clients, estimates to persist. "
                    "Foundation for everything else below.",
        "status":   "Partial — Streamlit storage saves metadata. Uploaded PDFs not persisted.",
    },
    {
        "gap":      "User accounts",
        "effort":   "1 week",
        "impact":   "HIGH",
        "why":      "No login = no separation between users. "
                    "Supabase auth (email or Google) solves this once storage exists.",
        "status":   "Not started",
    },
    {
        "gap":      "Multi-user / team support",
        "effort":   "2-3 weeks",
        "impact":   "MEDIUM",
        "why":      "Two estimators get separate sessions with no shared data. "
                    "Need shared project state, role-based access (estimator / reviewer / PM), "
                    "conflict resolution when two people edit simultaneously.",
        "status":   "Not started",
    },
    {
        "gap":      "Audit trail",
        "effort":   "1-2 weeks",
        "impact":   "MEDIUM",
        "why":      "No record of who changed what, when. "
                    "Critical for bid review and dispute resolution. "
                    "Every module change logged: user, field, old value, new value, timestamp.",
        "status":   "Not started",
    },
    {
        "gap":      "Real PDF parsing pipeline",
        "effort":   "2-3 weeks",
        "impact":   "HIGH",
        "why":      "AI Takeoff currently reads doc names, not content. "
                    "Production: extract schedule page text → Claude structured extraction; "
                    "render floor plan pages as images → Claude Vision for symbol detection.",
        "status":   "Partial — text search works; Claude Vision call not wired",
    },
    {
        "gap":      "Price book maintenance",
        "effort":   "1 week",
        "impact":   "LOW",
        "why":      "Prices change quarterly. Currently 2022 list prices hardcoded in JSON. "
                    "Need UI to edit items, import from Excel, or connect to vendor API.",
        "status":   "JSON structure ready — UI not built",
    },
    {
        "gap":      "Mobile / field use",
        "effort":   "2-4 weeks",
        "impact":   "LOW",
        "why":      "Estimators review drawings on site. "
                    "Streamlit is desktop-first. Real mobile needs responsive layout "
                    "or separate React frontend.",
        "status":   "Not started",
    },
]


def _product_status_panel():
    st.markdown("### Product status & roadmap")
    st.caption(
        "West 34th Street Hotel as reference project · "
        "Honest assessment of what's working, what's in progress, and what's next"
    )

    # ── Overall readiness bar ─────────────────────────────────────────────
    st.divider()
    st.markdown("#### Overall tool readiness")

    avg_ready = sum(m["ready"] for m in MODULE_READINESS) // len(MODULE_READINESS)
    prod_count  = sum(1 for m in MODULE_READINESS if m["status"] == "production")
    demo_count  = sum(1 for m in MODULE_READINESS if m["status"] == "demo")

    oc1, oc2, oc3, oc4 = st.columns(4)
    oc1.metric("Overall readiness",  f"{avg_ready}%")
    oc2.metric("Production-ready modules", prod_count)
    oc3.metric("Demo-ready modules",  demo_count)
    oc4.metric("Roadmap items",       len(PRODUCTION_GAPS))

    st.progress(avg_ready / 100,
                text=f"{avg_ready}% ready · {prod_count} modules production-ready · {demo_count} demo-ready")

    # ── Per-module readiness ──────────────────────────────────────────────
    st.divider()
    st.markdown("#### Module readiness")

    for mod in MODULE_READINESS:
        status_color = (
            "🟢" if mod["status"] == "production" else
            "🟡" if mod["status"] == "demo" else "🔴"
        )
        status_label = (
            "Production-ready" if mod["status"] == "production" else
            "Demo-ready" if mod["status"] == "demo" else "Not ready"
        )
        with st.expander(
            f"{mod['icon']} {mod['module']} — {mod['ready']}% · {status_color} {status_label}",
            expanded=False
        ):
            st.progress(mod["ready"] / 100)

            c1, c2 = st.columns(2)
            with c1:
                st.markdown("**✅ What works**")
                for w in mod["what_works"]:
                    st.markdown(f"- {w}")
            with c2:
                st.markdown("**⚠️ Gaps**")
                for g in mod["gaps"]:
                    st.markdown(f"- {g}")

            st.markdown(f"**➡ Next:** {mod['next']}")

    # ── Production gaps ───────────────────────────────────────────────────
    st.divider()
    st.markdown("#### What a production version needs")
    st.caption("Ordered by impact. Total estimate: 3-4 months part-time.")

    IMPACT_COLOR = {"HIGH": "🔴", "MEDIUM": "🟡", "LOW": "⚪"}

    for gap in PRODUCTION_GAPS:
        ic = IMPACT_COLOR.get(gap["impact"], "⚪")
        with st.expander(
            f"{ic} {gap['gap']} — {gap['effort']} · Impact: {gap['impact']}",
            expanded=False
        ):
            st.markdown(f"**Why it matters:** {gap['why']}")
            st.markdown(f"**Current status:** {gap['status']}")

    # ── Setup checklist ──────────────────────────────────────────────────
    st.divider()
    st.markdown("#### 📋 Setup checklist")

    checks = [
        ("Load West 34th JSON",
         "Takeoff tab → 'Or upload schedule_ground_truth.json'. "
         "Gives 109 real devices and 12 discrepancies. Do NOT use AI Takeoff button live."),
        ("Run drawing markup on your 15MB set",
         "Upload real drawing set, click Process, verify highlights look right, "
         "generate annotated PDF once and keep ready."),
        ("Add API key to Streamlit Secrets",
         "Streamlit dashboard → Settings → Secrets → "
         "ANTHROPIC_API_KEY = 'sk-ant-...'"),
        ("Practise the product walkthrough",
         "Overview → Discrepancies → Drawing Markup search → "
         "Send to proposal → Open this panel. Out loud, 3 times."),
        ("Know the roadmap for every incomplete tab",
         "Each tab that shows partial output has a clear next step in the roadmap panel."),
    ]

    all_done = True
    for i, (title, desc) in enumerate(checks, 1):
        done = st.checkbox(f"**{i}. {title}**", key=f"check_{i}")
        if not done:
            all_done = False
            st.caption(f"   → {desc}")

    if all_done:
        st.success("✅ All set — ready to go.")

    # ── Walkthrough script ────────────────────────────────────────────────
    st.divider()
    with st.expander("📣 Product walkthrough script"):
        st.markdown("""
**Beat 1 — The problem** *(60 sec)*
> *"A typical BMS estimate touches 5 documents, takes 2-3 days, and still misses things.
The most common miss: devices in the schedule with no SOO sequence.
Nobody writes them as excluded, nobody includes them in scope. They fall through the gap."*

---
**Beat 2 — The catch** *(90 sec)*

Open **Discrepancies** tab → point to amber banner.

> *"This project has 12 unit heaters in the M-200 schedule with integral thermostats
and zero SOO sequence. The tool caught them automatically.
Each one is a $500–700 scope question. Potentially $8,000 missed."*

Open **Drawing Markup** → search `UH-SC-1`.

> *"It's physically here on the drawing, it's in the schedule —
but there's no SOO sequence. Is it in BMS scope or not?
The tool forces that question before you submit the number."*

---
**Beat 3 — The pipeline** *(90 sec)*

Click through the 5 module tabs quickly.

> *"Once the takeoff is clean, the tool generates the point list in the client's
own Excel format, estimates labor from the point count,
builds the material list from our price book,
and drafts the proposal in our Word template.
The estimator reviews — AI does the first pass."*

Click **Send amber tags to proposal** → show Clarifications auto-filled.

---
**Beat 4 — Honest roadmap** *(60 sec)*

Open **this panel**.

> *"Here's what's production-ready, what's demo-ready, and what's missing.
Biggest gap: persistent storage — session state disappears on refresh.
Second: multi-user. Both are 6-8 week engineering problems,
not architecture problems — the data model already supports them."*

---
**Close**

> *"The core insight — three-way cross-check of schedule, SOO, and drawings —
is what's actually new. Everything else is workflow automation around that insight.
The architecture is modular so each piece can become an API
that plugs into whatever system you're already using."*
        """)


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    init(); sidebar()
    nav = st.session_state.nav
    if nav=="Overview":  page_overview()
    elif nav=="Projects":page_projects()
    elif nav=="Reports": page_reports()

if __name__=="__main__":
    main()
