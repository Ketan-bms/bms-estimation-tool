"""
output_generators.py - PRODUCTION MVP
Generates Word proposals and Excel estimates from analysis results
"""

import json
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from datetime import datetime


class OutputGenerator:
    """Generates Word proposals and Excel estimates"""
    
    def __init__(self, template_docx_path=None):
        self.template_docx_path = template_docx_path
        self.timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    
    # ============================================================================
    # WORD PROPOSAL GENERATION
    # ============================================================================
    
    @staticmethod
    def _num(value, default=0):
        """Coerce a model-supplied value to a number.

        Claude returns JSON, but a field typed as a number in the prompt can
        still come back as "40", "$150", "40 hours", or null. Every one of
        those blows up an f-string format spec like :,.0f, so coerce here
        rather than trusting the shape.
        """
        if isinstance(value, (int, float)):
            return value
        if value is None:
            return default
        cleaned = re.sub(r"[^0-9.\-]", "", str(value))
        try:
            return float(cleaned) if cleaned not in ("", "-", ".") else default
        except ValueError:
            return default

    @staticmethod
    def _safe_table_style(table, style_name):
        """Apply a table style, falling back if the template lacks it.

        A user-supplied .docx template may not define the style, in which
        case python-docx raises KeyError and kills the whole document.
        """
        try:
            table.style = style_name
        except KeyError:
            try:
                table.style = "Table Grid"
            except KeyError:
                pass

    # Common BMS/HVAC acronyms seen in real SOO headings, preserved as-is
    # when normalizing an ALL-CAPS heading. A length-based heuristic (e.g.
    # "5 letters or fewer is an acronym") looks appealing but misfires
    # constantly - WATER, VALVE, UNIT, COIL, FIRE are ordinary 4-5 letter
    # words, not acronyms, and would incorrectly stay uppercase. A curated
    # whitelist is slower to extend but doesn't have that failure mode.
    _KNOWN_ACRONYMS = {
        "DDC", "VFD", "AHU", "ERU", "DOAS", "HVAC", "BMS", "CT", "PCW",
        "CW", "HW", "CHW", "VAV", "FCU", "EF", "SF", "RF", "OA", "RA",
        "SA", "EA", "CO2", "UPS", "ATS", "FSD", "WSHP", "ASHP", "DX",
        "RTU", "MAU", "ERV", "HRV", "BACNET",
    }

    @classmethod
    def _normalize_title_case(cls, text):
        """If a heading is entirely upper-case, as some SOO formats write
        their section titles, convert it to a more readable case for a
        client-facing document - known acronyms are preserved, everything
        else is capitalized normally. Text that already has mixed case
        (meaning the source document formatted it the way it should read)
        is left completely untouched, since two different SOOs in the
        same proposal run can use different heading conventions and only
        one of them needs fixing.
        """
        if not text or text != text.upper():
            return text

        def fix_word(m):
            w = m.group(0)
            return w if w in cls._KNOWN_ACRONYMS else (w[0] + w[1:].lower())

        return re.sub(r'[A-Za-z]+', fix_word, text)

    @classmethod
    def _clean_section_title(cls, section_label):
        """Strip the SOO's own section numbering/lettering from a section
        label for display in the proposal, e.g. "1.8 PRIMARY CONDENSER
        WATER SYSTEM" -> "Primary Condenser Water System", or a
        subsectioned label like "3.2 SEQUENCE OF OPERATION - A. Energy
        Recovery Unit" -> "Energy Recovery Unit" (the innermost, most
        specific segment).

        Without this, a numbered proposal heading like "1. A. Energy
        Recovery Unit..." shows two competing numbering schemes stacked on
        top of each other - the SOO's own lettering leaking through next
        to this document's own numbering. Case is also normalized: two
        SOOs in the same run can use different heading conventions (one
        Title Case, one ALL CAPS), and headings should read consistently
        regardless of which source produced them.
        """
        name = section_label.rsplit(" - ", 1)[-1]
        name = re.sub(r'^\s*(\d+\.\d+|\d+|[A-Z]{1,2})\.?\s+', '', name)
        name = re.sub(r'\s*\(part \d+/\d+\)\s*$', '', name)
        name = re.sub(r'\s+', ' ', name).strip() or section_label
        return cls._normalize_title_case(name)

    def generate_word_proposal(self, analysis_results, project_name, output_path):
        """Generate a system-wise scope-of-work summary.

        Deliberately content-only: no letterhead, salutation, standard
        company boilerplate, or closing signature. This is meant to be
        pasted into an existing client-facing proposal template that
        already has all of that, not to stand alone as a formatted letter.
        Grouped by SOO system rather than a flat point table, mirroring how
        an estimator actually reads a scope section: one block per system,
        sub-grouped by equipment tag, with plain bulleted point names. No
        bid-letter metadata (bid number, drawing references, pricing) is
        included either, since none of that data exists in this pipeline -
        a placeholder or fabricated value would be worse than leaving it out.
        """
        try:
            if self.template_docx_path:
                doc = Document(self.template_docx_path)
            else:
                doc = Document()

            scope = analysis_results.get("scope", {})
            labor = analysis_results.get("labor_estimate", {})
            rfis = analysis_results.get("rfis", {})
            points = analysis_results.get("point_list", [])
            section_narratives = analysis_results.get("section_narratives", {})
            metadata = analysis_results.get("metadata", {})

            # ===== TITLE =====
            # No letterhead, salutation, or standard company boilerplate -
            # this is meant to be pasted directly into an existing client
            # proposal template that already has all of that, not to stand
            # alone as a formatted letter.
            title = doc.add_paragraph()
            title_run = title.add_run(f"{project_name} - Scope Summary")
            title_run.font.size = Pt(16)
            title_run.font.bold = True
            doc.add_paragraph()

            # ===== OVERVIEW =====
            n_systems = len({p.get("Source_Section", "") for p in points
                             if p.get("Source_Section")})
            overview = scope.get("project_overview", "")
            if overview:
                doc.add_paragraph(overview)
            doc.add_paragraph(
                f"This scope covers {n_systems} system(s) and "
                f"{len(points)} control point(s), based on review of the "
                f"provided Sequence of Operations."
            )

            systems_in_scope = scope.get("systems_in_scope", [])
            if systems_in_scope:
                doc.add_heading("Systems Covered", level=2)
                for system in systems_in_scope:
                    doc.add_paragraph(system, style='List Bullet')

            # ===== SCOPE OF WORK BY SYSTEM =====
            doc.add_heading("Scope of Work by System", level=1)

            if points:
                # Group points by their source SOO section, preserving the
                # order sections first appear in the point list (which
                # follows document order, since extraction runs
                # section-by-section top to bottom).
                sections_order = []
                by_section = {}
                for pt in points:
                    section = pt.get("Source_Section", "") or "Unlabeled section"
                    if section not in by_section:
                        by_section[section] = []
                        sections_order.append(section)
                    by_section[section].append(pt)

                for i, section in enumerate(sections_order, 1):
                    section_points = by_section[section]
                    pages = section_points[0].get("Source_Pages", "")
                    clean_title = self._clean_section_title(section)

                    heading = doc.add_paragraph()
                    heading_run = heading.add_run(
                        f"{i}. {clean_title}" + (f"  (SOO p{pages})" if pages else "")
                    )
                    heading_run.font.bold = True
                    heading_run.font.size = Pt(12)

                    # Narrative scope sentences pulled from the section's
                    # own text - "Furnish...", "Provide..." - read before
                    # the point bullets, matching how a real scope section
                    # actually presents: prose describing what will be
                    # done, then the specific points as a supporting list.
                    for sentence in section_narratives.get(section, []):
                        doc.add_paragraph(sentence)

                    # Show every extracted point - no confidence filtering
                    # here. Confidence grading is a review aid for the
                    # Points tab and Excel export, where a human checks the
                    # work before it goes anywhere; a real completed
                    # proposal states scope plainly, the same way every
                    # real example reviewed for this format does, with no
                    # hedging language in the delivered document.
                    equip_order = []
                    by_equip = {}
                    for pt in section_points:
                        raw_tag = str(pt.get("Equipment", "")).strip()
                        tag = raw_tag or "General"
                        if tag not in by_equip:
                            by_equip[tag] = []
                            equip_order.append(tag)
                        by_equip[tag].append(pt)

                    # Real proposals list every equipment tag for a system
                    # on one line - "Energy Recovery Units: ERU-08-01,
                    # ERU-08-02, ..." - even when there is only one tag,
                    # followed by ONE shared point list when every tag
                    # carries the same points (multiple instances of one
                    # equipment type). Tags are only split into separate
                    # sub-lists when they genuinely have different points -
                    # a section covering two distinct equipment types, for
                    # example a pump and a heat exchanger. A genuinely
                    # tagless system (no equipment named anywhere in that
                    # SOO section) gets no fabricated "General:" heading -
                    # none of the real documents this was checked against
                    # label an untagged system that way, they just move
                    # straight from narrative to the point list.
                    point_sets = {
                        tag: frozenset(p.get("Point_Name", "") for p in pts)
                        for tag, pts in by_equip.items()
                    }
                    same_points_throughout = len(set(point_sets.values())) <= 1
                    only_tagless = equip_order == ["General"]

                    if equip_order and same_points_throughout:
                        if not only_tagless:
                            tag_line = doc.add_paragraph()
                            tag_line.add_run(", ".join(equip_order)).font.bold = True
                        doc.add_paragraph("We will provide the following hardwired points:")
                        for pt in by_equip[equip_order[0]]:
                            name = pt.get("Point_Name", "")
                            qty = pt.get("Qty", "")
                            label = f"{name} (Qty. {qty})" if str(qty) not in ("", "1") else name
                            doc.add_paragraph(label, style='List Bullet')
                    else:
                        for tag in equip_order:
                            if tag != "General":
                                tag_para = doc.add_paragraph()
                                tag_run = tag_para.add_run(tag)
                                tag_run.font.bold = True
                            doc.add_paragraph("We will provide the following hardwired points:")
                            for pt in by_equip[tag]:
                                name = pt.get("Point_Name", "")
                                qty = pt.get("Qty", "")
                                label = f"{name} (Qty. {qty})" if str(qty) not in ("", "1") else name
                                doc.add_paragraph(label, style='List Bullet')

                    doc.add_paragraph()
            else:
                doc.add_paragraph("No points were extracted.")

            # ===== EXCLUSIONS =====
            exclusions = rfis.get("exclusions", [])
            if exclusions:
                doc.add_heading("Exclusions", level=1)
                for item in exclusions:
                    doc.add_paragraph(item, style='List Bullet')

            # ===== CLARIFICATIONS =====
            clarifications = rfis.get("rfis", [])
            if clarifications:
                doc.add_heading("Items Requiring Clarification", level=1)
                for item in clarifications:
                    doc.add_paragraph(item, style='List Bullet')

            # ===== LABOR (HOURS ONLY) =====
            role_totals = labor.get("role_totals", {})
            if role_totals:
                doc.add_heading("Estimated Labor (Hours)", level=1)

                table = doc.add_table(rows=1, cols=2)
                self._safe_table_style(table, 'Light Grid Accent 1')
                header_cells = table.rows[0].cells
                header_cells[0].text = "Role"
                header_cells[1].text = "Hours"
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                role_labels = {"tech": "Field Technician", "eng": "Engineering",
                               "soft": "Software", "gpc": "Graphics"}
                for key, label in role_labels.items():
                    row = table.add_row()
                    row.cells[0].text = label
                    row.cells[1].text = f"{self._num(role_totals.get(key)):,.1f}"

                total_row = table.add_row()
                total_row.cells[0].text = "TOTAL HOURS"
                total_row.cells[1].text = f"{self._num(labor.get('total_hours')):,.1f}"
                for cell in total_row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                if labor.get("assumptions"):
                    doc.add_paragraph()
                    assumptions_para = doc.add_paragraph()
                    assumptions_run = assumptions_para.add_run(
                        f"Assumptions: {labor['assumptions']}"
                    )
                    assumptions_run.font.italic = True
                    assumptions_run.font.size = Pt(9)

            doc.save(output_path)
            return True

        except Exception as e:
            raise RuntimeError("Word proposal generation failed: %s" % e) from e
    
    # ============================================================================
    # EXCEL ESTIMATE GENERATION
    # ============================================================================
    
    def generate_excel_estimate(self, analysis_results, project_name, output_path):
        """Generate Excel estimate matching your template structure"""
        
        try:
            wb = openpyxl.Workbook()
            
            # ===== SUMMARY SHEET =====
            ws_summary = wb.active
            ws_summary.title = "Summary"
            
            # Header
            ws_summary['A1'] = "General Project Information"
            ws_summary['B1'] = "Region:"
            ws_summary['C1'] = "New York"
            
            ws_summary['A2'] = "Project Name"
            ws_summary['C2'] = project_name
            
            ws_summary['A3'] = "Date"
            ws_summary['C3'] = self.timestamp
            
            # Material Summary
            ws_summary['A8'] = "Material Summary"
            ws_summary['A9'] = "Material cost:"
            ws_summary['A10'] = "Tax"
            ws_summary['A11'] = "Misc"
            ws_summary['A12'] = "Shipping"
            ws_summary['A13'] = "Warranty"
            
            labor = analysis_results.get("labor_estimate", {})
            ws_summary['A15'] = "Labor Summary"
            ws_summary['A16'] = "Total Labor Hours"
            ws_summary['C16'] = self._num(labor.get("total_hours"))
            
            # Format as currency (material summary rows only - labor is
            # hours-only, no dollar figure is produced for it)
            for row in [9, 10, 11, 12, 13]:
                ws_summary[f'C{row}'].number_format = '$#,##0.00'
            
            # ===== POINTS LIST SHEET =====
            ws_points = wb.create_sheet("Points List")
            
            # Headers
            # Field order matches the 11-column working format, with the
            # audit-trail columns appended so every row can be traced back
            # to the section and wording it came from.
            headers = [
                ("Panel", "Panel"),
                ("Equipment", "Equipment"),
                ("Point_Name", "Point_Name"),
                ("Control Device", "Control Device"),
                ("AI", "AI"), ("BI", "BI"), ("AO", "AO"), ("BO", "BO"),
                ("Qty", "Qty"),
                ("Description", "Description"),
                ("Confidence", "Confidence"),
                ("Source_Section", "Source Section"),
                ("Source_Pages", "SOO Pages"),
                ("Evidence", "Evidence (verbatim)"),
                ("Repeats_In_Sections", "Repeats"),
            ]
            for col_idx, (_, label) in enumerate(headers, 1):
                cell = ws_points.cell(1, col_idx)
                cell.value = label
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")

            points = analysis_results.get("point_list", [])
            conf_fill = {
                "high": PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid"),
                "medium": PatternFill(start_color="FFF7E0", end_color="FFF7E0", fill_type="solid"),
                "low": PatternFill(start_color="FCE8E6", end_color="FCE8E6", fill_type="solid"),
            }

            for row_idx, point in enumerate(points, 2):
                for col_idx, (key, _) in enumerate(headers, 1):
                    ws_points.cell(row_idx, col_idx).value = point.get(key, "")
                # Shade the confidence cell so low-confidence rows are
                # obvious when the sheet is reviewed by hand.
                fill = conf_fill.get(str(point.get("Confidence", "")).lower())
                if fill:
                    ws_points.cell(row_idx, 11).fill = fill

            ws_points.freeze_panes = "A2"
            
            # Adjust column widths
            for col, width in (("K", 12), ("L", 42), ("M", 11), ("N", 46), ("O", 9)):
                ws_points.column_dimensions[col].width = width
            ws_points.column_dimensions['A'].width = 12
            ws_points.column_dimensions['B'].width = 12
            ws_points.column_dimensions['C'].width = 25
            ws_points.column_dimensions['D'].width = 15
            
            # ===== LABOR (ES LABOR) SHEET =====
            # Three-row header matching the real estimating format: a group
            # label, the role each column bills to, then the actual field
            # name. Totals are live Excel formulas, not values computed in
            # Python - editing Quantity in the sheet recalculates every
            # total automatically, which is the whole point of building it
            # this way rather than writing static numbers.
            ws_labor = wb.create_sheet("Labor Estimate")

            HOUR_COLS = ["PanelFabUnitHrs", "EngOrigUnitHrs", "EngCopyUnitHrs",
                        "SoftOrigUnitHrs", "SoftCopyUnitHrs", "ScreenOrigUnitHrs",
                        "ScreenCopyUnitHrs", "StartupUnitHrs", "CommissUnitHrs"]
            ROLE_ROW = ["Tech", "Eng", "Eng", "Sfw", "Sfw", "Gpc", "Gpc", "Tech", "Tech"]
            # Hour columns start at column C (A=Panel, B=Qty)
            first_hour_col = 3
            last_hour_col = first_hour_col + len(HOUR_COLS) - 1  # column K

            header_fill = PatternFill(start_color="1F3864", end_color="1F3864", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")
            role_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")

            ws_labor.merge_cells(
                start_row=1, start_column=first_hour_col, end_row=1, end_column=last_hour_col
            )
            es_cell = ws_labor.cell(1, first_hour_col)
            es_cell.value = "ES LABOR"
            es_cell.font = header_font
            es_cell.fill = header_fill
            es_cell.alignment = Alignment(horizontal="center")

            for i, role in enumerate(ROLE_ROW):
                cell = ws_labor.cell(2, first_hour_col + i)
                cell.value = role
                cell.font = Font(bold=True)
                cell.fill = role_fill
                cell.alignment = Alignment(horizontal="center")

            headers_row3 = ["Panel", "Qty"] + HOUR_COLS + [
                "TechTotalHrs", "EngTotalHrs", "SoftTotalHrs", "ScreenTotalHrs", "SystemTotalHrs"
            ]
            for i, h in enumerate(headers_row3, 1):
                cell = ws_labor.cell(3, i)
                cell.value = h
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center", wrap_text=True)

            systems = labor.get("systems", [])
            row = 4
            for sys_row in systems:
                ws_labor.cell(row, 1).value = sys_row.get("panel", "")
                ws_labor.cell(row, 2).value = self._num(sys_row.get("quantity"), default=1)
                for i, key in enumerate(["panel_fab_hours", "eng_orig_hours", "eng_copy_hours",
                                         "soft_orig_hours", "soft_copy_hours", "screen_orig_hours",
                                         "screen_copy_hours", "startup_hours", "commiss_hours"]):
                    ws_labor.cell(row, first_hour_col + i).value = self._num(sys_row.get(key))

                # Live formulas: PanelFab/Startup/Commiss scale directly with
                # Qty; Eng/Soft/Screen apply Copy hours only to units beyond
                # the first (MAX(Qty-1,0) so a Qty of 0 or 1 never goes
                # negative). Editing the Qty cell recalculates all of these.
                ws_labor.cell(row, 12).value = f"=(C{row}+J{row}+K{row})*B{row}"       # Tech
                ws_labor.cell(row, 13).value = f"=D{row}+E{row}*MAX(B{row}-1,0)"       # Eng
                ws_labor.cell(row, 14).value = f"=F{row}+G{row}*MAX(B{row}-1,0)"       # Soft
                ws_labor.cell(row, 15).value = f"=H{row}+I{row}*MAX(B{row}-1,0)"       # Screen
                ws_labor.cell(row, 16).value = f"=L{row}+M{row}+N{row}+O{row}"         # System total
                row += 1

            last_data_row = row - 1
            if systems:
                total_row = row
                ws_labor.cell(total_row, 1).value = "TOTAL"
                ws_labor.cell(total_row, 1).font = Font(bold=True)
                for col in range(12, 17):
                    letter = get_column_letter(col)
                    cell = ws_labor.cell(total_row, col)
                    cell.value = f"=SUM({letter}4:{letter}{last_data_row})"
                    cell.font = Font(bold=True)

            for col, width in [("A", 34), ("B", 6)] + [
                (get_column_letter(c), 12) for c in range(first_hour_col, 17)
            ]:
                ws_labor.column_dimensions[col].width = width
            ws_labor.freeze_panes = "C4"

            if labor.get("assumptions"):
                notes_row = row + 2
                ws_labor.cell(notes_row, 1).value = f"Assumptions: {labor['assumptions']}"
                ws_labor.cell(notes_row, 1).font = Font(italic=True, size=9)

            # ===== ANALYSIS NOTES SHEET =====
            ws_notes = wb.create_sheet("Analysis Notes")
            
            ws_notes['A1'] = "Scope Analysis Notes"
            scope = analysis_results.get("scope", {})
            ws_notes['A3'] = "Systems in Scope:"
            row = 4
            for system in scope.get("systems_in_scope", []):
                ws_notes[f'A{row}'] = system
                row += 1
            
            rfis = analysis_results.get("rfis", {})
            ws_notes['A10'] = "RFIs (Items needing clarification):"
            row = 11
            for rfi in rfis.get("rfis", []):
                ws_notes[f'A{row}'] = rfi
                row += 1
            
            ws_notes['A20'] = "Exclusions:"
            row = 21
            for exc in rfis.get("exclusions", []):
                ws_notes[f'A{row}'] = exc
                row += 1
            
            # Column widths
            ws_notes.column_dimensions['A'].width = 60
            
            # Save workbook
            wb.save(output_path)
            return True
            
        except Exception as e:
            raise RuntimeError("Excel estimate generation failed: %s" % e) from e
    
    # ============================================================================
    # BULK EXPORT
    # ============================================================================
    
    # ============================================================================
    # STANDALONE POINT LIST EXPORT
    # ============================================================================

    POINT_COLUMNS = [
        ("Panel", "Panel", 14),
        ("Equipment", "Equipment", 16),
        ("Point_Name", "Point Name", 30),
        ("Control Device", "Control Device", 16),
        ("AI", "AI", 5),
        ("BI", "BI", 5),
        ("AO", "AO", 5),
        ("BO", "BO", 5),
        ("Qty", "Qty", 6),
        ("Description", "Remarks", 46),
        ("Confidence", "Confidence", 12),
        ("Source_Section", "Source Section", 42),
        ("Source_Pages", "SOO Pages", 11),
        ("Evidence", "Evidence (verbatim)", 46),
        ("Repeats_In_Sections", "Repeats", 9),
    ]

    def generate_point_list_excel(self, analysis_results, project_name, output_path):
        """Write the point list as a standalone workbook.

        Two sheets: the points themselves, and a per-section summary so a
        reviewer can see where each block of points came from and which
        sections produced none. Filtering and freezing are set up so the
        sheet is usable for review without further formatting.
        """
        try:
            points = analysis_results.get("point_list", [])
            metadata = analysis_results.get("metadata", {})

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Point List"

            header_fill = PatternFill(start_color="1F3864", end_color="1F3864",
                                      fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")

            for col_idx, (_, label, width) in enumerate(self.POINT_COLUMNS, 1):
                cell = ws.cell(1, col_idx)
                cell.value = label
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center", vertical="center",
                                           wrap_text=True)
                ws.column_dimensions[get_column_letter(col_idx)].width = width

            conf_fill = {
                "high": PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid"),
                "medium": PatternFill(start_color="FFF7E0", end_color="FFF7E0", fill_type="solid"),
                "low": PatternFill(start_color="FCE8E6", end_color="FCE8E6", fill_type="solid"),
            }
            conf_col = [k for k, _, _ in self.POINT_COLUMNS].index("Confidence") + 1

            for row_idx, point in enumerate(points, 2):
                for col_idx, (key, _, _) in enumerate(self.POINT_COLUMNS, 1):
                    value = point.get(key, "")
                    if key == "Qty":
                        value = self._num(value, default=1)
                    ws.cell(row_idx, col_idx).value = value
                    ws.cell(row_idx, col_idx).alignment = Alignment(
                        vertical="top", wrap_text=(key in ("Description", "Evidence"))
                    )
                fill = conf_fill.get(str(point.get("Confidence", "")).lower())
                if fill:
                    ws.cell(row_idx, conf_col).fill = fill

            ws.freeze_panes = "C2"
            if points:
                ws.auto_filter.ref = (
                    f"A1:{get_column_letter(len(self.POINT_COLUMNS))}{len(points) + 1}"
                )

            # ---- Summary sheet ----
            ws2 = wb.create_sheet("Summary")
            ws2["A1"] = "Point List Summary"
            ws2["A1"].font = Font(bold=True, size=14)

            ws2["A3"] = "Project"
            ws2["B3"] = project_name
            ws2["A4"] = "Generated"
            ws2["B4"] = self.timestamp
            ws2["A5"] = "SOO pages"
            ws2["B5"] = metadata.get("soo_pages", "")
            ws2["A6"] = "Rows in list"
            ws2["B6"] = len(points)
            ws2["A7"] = "Total I/O (sum of Qty)"
            ws2["B7"] = metadata.get("total_i_o_count", "")

            coverage = metadata.get("coverage", {}) or {}
            ws2["A8"] = "Document covered"
            ws2["B8"] = f"{coverage.get('coverage_pct', '')}%"

            counts = metadata.get("confidence_counts", {}) or {}
            ws2["A10"] = "Confidence"
            ws2["A10"].font = Font(bold=True)
            for i, level in enumerate(("high", "medium", "low"), start=11):
                ws2[f"A{i}"] = level.title()
                ws2[f"B{i}"] = counts.get(level, 0)

            io_totals = {}
            for point in points:
                for io in ("AI", "BI", "AO", "BO"):
                    if str(point.get(io, "")).strip():
                        io_totals[io] = io_totals.get(io, 0) + self._num(
                            point.get("Qty"), default=1
                        )
            ws2["A15"] = "I/O by type (Qty weighted)"
            ws2["A15"].font = Font(bold=True)
            for i, io in enumerate(("AI", "BI", "AO", "BO"), start=16):
                ws2[f"A{i}"] = io
                ws2[f"B{i}"] = io_totals.get(io, 0)

            ws2["A21"] = "Points by section"
            ws2["A21"].font = Font(bold=True)
            ws2["A22"] = "Section"
            ws2["B22"] = "Pages"
            ws2["C22"] = "Points"
            ws2["D22"] = "Status"
            for c in ("A22", "B22", "C22", "D22"):
                ws2[c].font = Font(bold=True)

            for i, row in enumerate(metadata.get("sections", []), start=23):
                ws2[f"A{i}"] = row.get("section", "")
                ws2[f"B{i}"] = row.get("pages", "")
                ws2[f"C{i}"] = row.get("points", 0)
                ws2[f"D{i}"] = row.get("status", "")

            ws2.column_dimensions["A"].width = 46
            ws2.column_dimensions["B"].width = 14
            ws2.column_dimensions["C"].width = 10
            ws2.column_dimensions["D"].width = 10

            wb.save(output_path)
            return True

        except Exception as e:
            raise RuntimeError("Point list export failed: %s" % e) from e

    def export_all_outputs(self, analysis_results, project_name, output_dir):
        """Generate all output files at once"""
        
        word_path = f"{output_dir}/{project_name}_Proposal.docx"
        excel_path = f"{output_dir}/{project_name}_Estimate.xlsx"
        
        self.generate_word_proposal(analysis_results, project_name, word_path)
        self.generate_excel_estimate(analysis_results, project_name, excel_path)

        outputs = {}
        if os.path.exists(word_path):
            outputs["proposal"] = word_path
        if os.path.exists(excel_path):
            outputs["estimate"] = excel_path
        return outputs


# ============================================================================
# EXAMPLE USAGE
# ============================================================================

if __name__ == "__main__":
    # Load analysis results from JSON
    with open("analysis_results.json", "r") as f:
        results = json.load(f)
    
    # Generate outputs
    generator = OutputGenerator(template_docx_path="proposal_template.docx")
    
    generator.export_all_outputs(
        analysis_results=results,
        project_name="Example_Project",
        output_dir="./outputs"
    )
