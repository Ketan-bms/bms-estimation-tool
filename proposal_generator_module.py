"""
proposal_generator_module.py
Independent proposal generator where user provides template, I follow format.
Keeps "Qty x" notation as specified.
"""

import json
from io import BytesIO


def analyze_proposal_template(docx_bytes):
    """
    Analyze user's proposal template to extract format/structure.
    Returns: Template metadata (sections, formatting, example)
    """
    try:
        from docx import Document
        doc = Document(BytesIO(docx_bytes))
        
        metadata = {
            "filename": "proposal_template",
            "num_pages": len(doc.paragraphs),
            "num_tables": len(doc.tables),
            "sections": [],
            "format_notes": []
        }
        
        # Extract section headings
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                metadata["sections"].append(para.text)
        
        # Extract first table structure if exists
        if doc.tables:
            first_table = doc.tables[0]
            first_row = [cell.text for cell in first_table.rows[0].cells]
            metadata["table_columns"] = first_row
        
        return metadata
    except Exception as e:
        return {"error": str(e), "fallback": "Use standard TEC Building Systems format"}


def generate_proposal_prompt(project_name, soo_text, template_analysis, client_name=""):
    """
    Generate Claude prompt to create proposal matching user's template format.
    Keeps "Qty x" notation.
    """
    
    template_guidance = ""
    if template_analysis and "sections" in template_analysis:
        template_guidance = f"""
TEMPLATE STRUCTURE TO FOLLOW:
Sections in order: {', '.join(template_analysis['sections'][:5])}

TABLE COLUMNS: {', '.join(template_analysis.get('table_columns', []))}
"""
    
    prompt = f"""You are a senior BMS proposal writer for TEC Building Systems.
Generate a complete project proposal following the structure below.

PROJECT: {project_name}
CLIENT: {client_name or "Unspecified"}

{template_guidance}

SEQUENCE OF OPERATIONS (use to extract scope):
{soo_text[:8000]}

OUTPUT RULES:
1. Start with "Date: [today]" and "TEC Bid No: [auto]"
2. Include all required sections:
   - Project Title & Documents Referenced
   - Scope of Work (organized by system, keep "Qty x" notation)
   - Equipment List (with Qty x format)
   - Pricing (placeholder: $ XX)
   - Notes & Clarifications
   - Exclusions (explicit list)
3. System organization (from SOO):
   - ASHP systems (with Qty x and specific controls)
   - DOAS systems (with Qty x and specific controls)
   - AHU systems (with Qty x and specific controls)
   - FCU systems (with Qty x and specific controls)
   - VAV boxes (with Qty x and specific controls)
   - Pump systems (with Qty x and specific controls)
   - Exhaust fans and specialty equipment
4. For each system, list:
   - Quantity (Qty x format, e.g., "Qty 3")
   - Specific control points (start/stop, status, sensors, alarms)
   - Integration method (DDC controller, BACnet, local control, etc.)
   - Wiring scope (field wiring, terminations, etc.)

5. Keep professional tone (matching TEC template if provided)
6. Use bullet points for clarity
7. Include standard warranty/training/commissioning language
8. End with signature line: "Sincerely, TEC Building Systems, LLC"

CRITICAL: Output as properly formatted text (not JSON or markdown).
Start with Date line, end with signature.
"""
    
    return prompt


def extract_proposal_structure(proposal_text):
    """
    Parse generated proposal and extract key sections.
    Returns structured data for further refinement.
    """
    sections = {
        "header": "",
        "scope_by_system": {},
        "equipment_list": [],
        "pricing": "",
        "notes": "",
        "exclusions": []
    }
    
    lines = proposal_text.split("\n")
    current_section = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Detect section headers
        if line.startswith("Scope of Work:") or line.lower().startswith("scope"):
            current_section = "scope"
        elif line.startswith("Pricing:") or line.lower().startswith("pricing"):
            current_section = "pricing"
        elif "Exclusions" in line:
            current_section = "exclusions"
        elif "Notes" in line or "Clarifications" in line:
            current_section = "notes"
        elif line.startswith("Date:"):
            sections["header"] = line
        
        # Collect content
        if current_section == "scope":
            if "Qty " in line or "qty " in line:
                sections["equipment_list"].append(line)
        elif current_section == "exclusions":
            if line.startswith("-") or line.startswith("•"):
                sections["exclusions"].append(line)
        elif current_section == "pricing":
            if "$" in line:
                sections["pricing"] += line + "\n"
    
    return sections


if __name__ == "__main__":
    print("Proposal Generator Module")
    print("=" * 70)
    
    # Example template analysis
    example_analysis = {
        "sections": ["Scope of Work", "Equipment List", "Labor", "Pricing", "Notes"],
        "table_columns": ["System", "Equipment", "Qty", "Description", "Control Points"]
    }
    
    prompt = generate_proposal_prompt(
        "Test Project",
        "ASHP-1 with 3x cooling, DHW, freeze protection...",
        example_analysis,
        "Test Client"
    )
    
    print("\nGenerated Prompt (first 500 chars):")
    print(prompt[:500])
    print("\n✅ Module ready for integration into app.py")
